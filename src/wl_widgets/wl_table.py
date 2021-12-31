#
# Wordless: Widgets - Table
#
# Copyright (C) 2018-2021  Ye Lei (叶磊)
#
# This source file is licensed under GNU GPLv3.
# For details, see: https://github.com/BLKSerene/Wordless/blob/master/LICENSE.txt
#
# All other rights reserved.
#

import csv
import os
import random
import re

from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *

import docx
import openpyxl

from wl_checking import wl_checking_misc
from wl_dialogs import wl_dialog_misc, wl_msg_box
from wl_text import wl_text_utils
from wl_utils import wl_misc, wl_threading
from wl_widgets import wl_box, wl_button, wl_msg

class Wl_Worker_Export_Table(wl_threading.Wl_Worker):
    worker_done = pyqtSignal(bool, str)

    def run(self):
        if 'headers_int' not in self.table.__dict__:
            self.table.headers_int = []
        if 'headers_float' not in self.table.__dict__:
            self.table.headers_float = []
        if 'headers_pct' not in self.table.__dict__:
            self.table.headers_pct = []

        settings_concordancer = self.main.settings_custom['concordancer']['zapping_settings']

        # Check file permissions
        try:
            if not self.rows_export:
                self.rows_export = list(range(self.table.rowCount()))

            file_path_src = re.sub(r'\.([a-z]+?)$', r'_source.\1', self.file_path)
            file_path_tgt = re.sub(r'\.([a-z]+?)$', r'_target.\1', self.file_path)

            len_rows = len(self.rows_export)

            # CSV files
            if self.file_type == self.tr('CSV File (*.csv)'):
                encoding = self.main.settings_custom['export']['tables']['default_encoding']

                # Concordancer
                if self.table.tab == 'concordancer':
                    with open(self.file_path, 'w', encoding = encoding, newline = '') as f:
                        csv_writer = csv.writer(f)

                        # Horizontal Headers
                        csv_writer.writerow([self.table.horizontalHeaderItem(col).text().strip()
                                             for col in range(self.table.columnCount())])

                        # Cells
                        for i, row in enumerate(self.rows_export):
                            row_to_export = []

                            for col in range(self.table.columnCount()):
                                if self.table.item(row, col):
                                    cell_text = self.table.item(row, col).text()
                                else:
                                    cell_text = self.table.cellWidget(row, col).text()
                                    cell_text = wl_text_utils.html_to_text(cell_text)

                                row_to_export.append(cell_text)

                            csv_writer.writerow(row_to_export)

                            self.progress_updated.emit(self.tr(f'Exporting table ... ({i + 1} / {len_rows})'))
                # Concordancer (Parallel Mode)
                elif self.table.tab == 'concordancer_parallel':
                    # Source file
                    with open(file_path_src, 'w', encoding = encoding, newline = '') as f:
                        csv_writer = csv.writer(f)

                        # Horizontal Headers
                        csv_writer.writerow([self.table.linked_tables[0].horizontalHeaderItem(col).text().strip()
                                             for col in range(self.table.linked_tables[0].columnCount())])

                        # Cells
                        for i, row in enumerate(self.rows_export):
                            row_to_export = []

                            for col in range(self.table.linked_tables[0].columnCount()):
                                if self.table.linked_tables[0].item(row, col):
                                    cell_text = self.table.linked_tables[0].item(row, col).text()
                                else:
                                    cell_text = self.table.linked_tables[0].cellWidget(row, col).text()
                                    cell_text = wl_text_utils.html_to_text(cell_text)

                                row_to_export.append(cell_text)

                            csv_writer.writerow(row_to_export)

                            self.progress_updated.emit(self.tr(f'Exporting table ... ({i + 1} / {len_rows * 2})'))

                    # Target file
                    with open(file_path_tgt, 'w', encoding = encoding, newline = '') as f:
                        csv_writer = csv.writer(f)

                        # Horizontal Headers
                        csv_writer.writerow([self.table.horizontalHeaderItem(col).text().strip()
                                             for col in range(self.table.columnCount())])

                        # Cells
                        for i, row in enumerate(self.rows_export):
                            row_to_export = []

                            for col in range(self.table.columnCount()):
                                if self.table.item(row, col):
                                    cell_text = self.table.item(row, col).text()
                                else:
                                    cell_text = self.table.cellWidget(row, col).text()
                                    cell_text = wl_text_utils.html_to_text(cell_text)

                                row_to_export.append(cell_text)

                            csv_writer.writerow(row_to_export)

                            self.progress_updated.emit(self.tr(f'Exporting table ... ({len_rows + i + 1} / {len_rows * 2})'))
                else:
                    with open(self.file_path, 'w', encoding = encoding, newline = '') as f:
                        csv_writer = csv.writer(f)

                        if self.table.header_orientation == 'horizontal':
                            # Horizontal Headers
                            csv_writer.writerow([self.table.horizontalHeaderItem(col).text().strip()
                                                 for col in range(self.table.columnCount())])

                            # Cells
                            for i, row in enumerate(self.rows_export):
                                row_to_export = []

                                for col in range(self.table.columnCount()):
                                    row_to_export.append(self.table.item(row, col).text().strip())

                                csv_writer.writerow(row_to_export)

                                self.progress_updated.emit(self.tr(f'Exporting table ... ({i + 1} / {len_rows})'))
                        else:
                            # Horizontal Headers
                            csv_writer.writerow([''] +
                                                [self.table.horizontalHeaderItem(col).text().strip()
                                                 for col in range(self.table.columnCount())])

                            # Vertical Headers & Cells
                            for i, row in enumerate(self.rows_export):
                                row_to_export = [self.table.verticalHeaderItem(row).text().strip()]

                                for col in range(self.table.columnCount()):
                                    row_to_export.append(self.table.item(row, col).text().strip())

                                csv_writer.writerow(row_to_export)

                                self.progress_updated.emit(self.tr(f'Exporting table ... ({i + 1} / {len_rows})'))
            # Excel workbooks
            elif self.file_type == self.tr('Excel Workbook (*.xlsx)'):
                dpi_horizontal = QApplication.primaryScreen().logicalDotsPerInchX()
                dpi_vertical = QApplication.primaryScreen().logicalDotsPerInchY()

                # Concordancer
                if self.table.tab == 'concordancer':
                    workbook = openpyxl.Workbook()
                    worksheet = workbook.active

                    worksheet.freeze_panes = 'A2'

                    dpi_horizontal = QApplication.primaryScreen().logicalDotsPerInchX()
                    dpi_vertical = QApplication.primaryScreen().logicalDotsPerInchY()

                    # Horizontal Headers
                    for col in range(self.table.columnCount()):
                        cell = worksheet.cell(1, 1 + col)
                        cell.value = self.table.horizontalHeaderItem(col).text()

                        self.style_header_horizontal(cell, self.table.horizontalHeaderItem(col))

                        worksheet.column_dimensions[openpyxl.utils.get_column_letter(1 + col)].width = self.table.horizontalHeader().sectionSize(col) / dpi_horizontal * 13 + 3

                    # Cells
                    for row_cell, row_item in enumerate(self.rows_export):
                        for col in range(self.table.columnCount()):
                            # Left
                            if col == 0:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_text_utils.html_to_text(self.table.cellWidget(row_item, col).text())
                                # Remove illegal characters
                                cell_val = re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', cell_val)
                                cell.value = cell_val

                                cell.font = openpyxl.styles.Font(
                                    name = self.table.cellWidget(row_item, col).font().family(),
                                    size = 8,
                                    color = '292929'
                                )
                                cell.alignment = openpyxl.styles.Alignment(
                                    horizontal = 'right',
                                    vertical = 'center'
                                )
                            # Node
                            elif col == 1:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_text_utils.html_to_text(self.table.cellWidget(row_item, col).text())
                                # Remove illegal characters
                                cell_val = re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', cell_val)
                                cell.value = cell_val

                                self.style_cell_text(cell, self.table.cellWidget(row_item, col))

                                cell.font = openpyxl.styles.Font(
                                    name = self.table.cellWidget(row_item, col).font().family(),
                                    size = 8,
                                    bold = True,
                                    color = 'FF0000'
                                )
                                cell.alignment = openpyxl.styles.Alignment(
                                    horizontal = 'center',
                                    vertical = 'center'
                                )
                            # Right
                            elif col == 2:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_text_utils.html_to_text(self.table.cellWidget(row_item, col).text())
                                # Remove illegal characters
                                cell_val = re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', cell_val)
                                cell.value = cell_val

                                cell.font = openpyxl.styles.Font(
                                    name = self.table.cellWidget(row_item, col).font().family(),
                                    size = 8,
                                    color = '292929'
                                )
                                cell.alignment = openpyxl.styles.Alignment(
                                    horizontal = 'left',
                                    vertical = 'center'
                                )
                            else:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = cell_text = self.table.item(row_item, col).text()
                                # Remove illegal characters
                                cell_val = re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', cell_val)
                                cell.value = cell_val

                                if (col in self.table.headers_int or
                                    col in self.table.headers_float or
                                    col in self.table.headers_pct):
                                    self.style_cell_num(cell, self.table.item(row_item, col))
                                else:
                                    self.style_cell_text(cell, self.table.item(row_item, col))

                            self.progress_updated.emit(self.tr(f'Exporting table ... ({row_cell + 1} / {len_rows})'))

                    # Row Height
                    worksheet.row_dimensions[1].height = self.table.horizontalHeader().height() / dpi_vertical * 72

                    for i, _ in enumerate(worksheet.rows):
                        worksheet.row_dimensions[2 + i].height = self.table.verticalHeader().sectionSize(0) / dpi_vertical * 72

                    self.progress_updated.emit(self.tr(f'Saving file ...'))

                    workbook.save(self.file_path)
                # Concordancer (Parallel Mode)
                elif self.table.tab == 'concordancer_parallel':
                    # Source file
                    workbook = openpyxl.Workbook()
                    worksheet = workbook.active

                    worksheet.freeze_panes = 'A2'

                    # Horizontal Headers
                    for col in range(self.table.linked_tables[0].columnCount()):
                        cell = worksheet.cell(1, 1 + col)
                        cell.value = self.table.linked_tables[0].horizontalHeaderItem(col).text()

                        self.style_header_horizontal(cell, self.table.linked_tables[0].horizontalHeaderItem(col))

                        worksheet.column_dimensions[openpyxl.utils.get_column_letter(1 + col)].width = self.table.linked_tables[0].horizontalHeader().sectionSize(col) / dpi_horizontal * 13 + 3

                    # Cells
                    for row_cell, row_item in enumerate(self.rows_export):
                        for col in range(self.table.linked_tables[0].columnCount()):
                            # Left
                            if col == 0:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_text_utils.html_to_text(self.table.linked_tables[0].cellWidget(row_item, col).text())
                                # Remove illegal characters
                                cell_val = re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', cell_val)
                                cell.value = cell_val

                                cell.font = openpyxl.styles.Font(
                                    name = self.table.linked_tables[0].cellWidget(row_item, col).font().family(),
                                    size = 8,
                                    color = '292929'
                                )
                                cell.alignment = openpyxl.styles.Alignment(
                                    horizontal = 'right',
                                    vertical = 'center'
                                )
                            # Node
                            elif col == 1:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_text_utils.html_to_text(self.table.linked_tables[0].cellWidget(row_item, col).text())
                                # Remove illegal characters
                                cell_val = re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', cell_val)
                                cell.value = cell_val

                                self.style_cell_text(cell, self.table.linked_tables[0].cellWidget(row_item, col))

                                cell.font = openpyxl.styles.Font(
                                    name = self.table.linked_tables[0].cellWidget(row_item, col).font().family(),
                                    size = 8,
                                    bold = True,
                                    color = 'FF0000'
                                )
                                cell.alignment = openpyxl.styles.Alignment(
                                    horizontal = 'center',
                                    vertical = 'center'
                                )
                            # Right
                            elif col == 2:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_text_utils.html_to_text(self.table.linked_tables[0].cellWidget(row_item, col).text())
                                # Remove illegal characters
                                cell_val = re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', cell_val)
                                cell.value = cell_val

                                cell.font = openpyxl.styles.Font(
                                    name = self.table.linked_tables[0].cellWidget(row_item, col).font().family(),
                                    size = 8,
                                    color = '292929'
                                )
                                cell.alignment = openpyxl.styles.Alignment(
                                    horizontal = 'left',
                                    vertical = 'center'
                                )
                            else:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = cell_text = self.table.linked_tables[0].item(row_item, col).text()
                                # Remove illegal characters
                                cell_val = re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', cell_val)
                                cell.value = cell_val

                                if (col in self.table.linked_tables[0].headers_int or
                                    col in self.table.linked_tables[0].headers_float or
                                    col in self.table.linked_tables[0].headers_pct):
                                    self.style_cell_num(cell, self.table.linked_tables[0].item(row_item, col))
                                else:
                                    self.style_cell_text(cell, self.table.linked_tables[0].item(row_item, col))

                            self.progress_updated.emit(self.tr(f'Exporting table ... ({row_cell + 1} / {len_rows * 2})'))

                    # Row Height
                    worksheet.row_dimensions[1].height = self.table.linked_tables[0].horizontalHeader().height() / dpi_vertical * 72

                    for i, _ in enumerate(worksheet.rows):
                        worksheet.row_dimensions[2 + i].height = self.table.linked_tables[0].verticalHeader().sectionSize(0) / dpi_vertical * 72

                    self.progress_updated.emit(self.tr(f'Saving source file ...'))

                    workbook.save(file_path_src)

                    # Source file
                    workbook = openpyxl.Workbook()
                    worksheet = workbook.active

                    worksheet.freeze_panes = 'A2'

                    # Horizontal Headers
                    for col in range(self.table.columnCount()):
                        cell = worksheet.cell(1, 1 + col)
                        cell.value = self.table.horizontalHeaderItem(col).text()

                        self.style_header_horizontal(cell, self.table.horizontalHeaderItem(col))

                        worksheet.column_dimensions[openpyxl.utils.get_column_letter(1 + col)].width = self.table.horizontalHeader().sectionSize(col) / dpi_horizontal * 13 + 3

                    # Cells
                    for row_cell, row_item in enumerate(self.rows_export):
                        for col in range(self.table.columnCount()):
                            # Parallel Text
                            if col == 0:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_text_utils.html_to_text(self.table.cellWidget(row_item, col).text())
                                # Remove illegal characters
                                cell_val = re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', cell_val)
                                cell.value = cell_val

                                self.style_cell_text(cell, self.table.cellWidget(row_item, col))

                                cell.font = openpyxl.styles.Font(
                                    name = self.table.cellWidget(row_item, col).font().family(),
                                    size = 8,
                                )
                                cell.alignment = openpyxl.styles.Alignment(
                                    horizontal = 'center',
                                    vertical = 'center'
                                )
                            else:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = cell_text = self.table.item(row_item, col).text()
                                # Remove illegal characters
                                cell_val = re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', cell_val)
                                cell.value = cell_val

                                if (col in self.table.headers_int or
                                    col in self.table.headers_float or
                                    col in self.table.headers_pct):
                                    self.style_cell_num(cell, self.table.item(row_item, col))
                                else:
                                    self.style_cell_text(cell, self.table.item(row_item, col))

                            self.progress_updated.emit(self.tr(f'Exporting table ... ({len_rows + row_cell + 1} / {len_rows * 2})'))

                    # Row Height
                    worksheet.row_dimensions[1].height = self.table.horizontalHeader().height() / dpi_vertical * 72

                    for i, _ in enumerate(worksheet.rows):
                        worksheet.row_dimensions[2 + i].height = self.table.verticalHeader().sectionSize(0) / dpi_vertical * 72

                    self.progress_updated.emit(self.tr(f'Saving target file ...'))

                    workbook.save(file_path_tgt)
                else:
                    workbook = openpyxl.Workbook()
                    worksheet = workbook.active

                    worksheet.freeze_panes = 'B2'

                    if self.table.header_orientation == 'horizontal':
                        # Horizontal Headers
                        for col in range(self.table.columnCount()):
                            cell = worksheet.cell(1, 1 + col)
                            cell.value = self.table.horizontalHeaderItem(col).text()

                            self.style_header_horizontal(cell, self.table.horizontalHeaderItem(col))

                            worksheet.column_dimensions[openpyxl.utils.get_column_letter(1 + col)].width = self.table.horizontalHeader().sectionSize(col) / dpi_horizontal * 13 + 3

                        # Cells
                        for row_cell, row_item in enumerate(self.rows_export):
                            for col in range(self.table.columnCount()):
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = self.table.item(row_item, col).text()
                                # Remove illegal characters
                                cell_val = re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', cell_val)
                                cell.value = cell_val

                                if (col in self.table.headers_int or
                                    col in self.table.headers_float or
                                    col in self.table.headers_pct):
                                    self.style_cell_num(cell, self.table.item(row_item, col))
                                else:
                                    self.style_cell_text(cell, self.table.item(row_item, col))

                            self.progress_updated.emit(self.tr(f'Exporting table ... ({row_cell + 1} / {len_rows})'))
                    else:
                        # Horizontal Headers
                        for col in range(self.table.columnCount()):
                            cell = worksheet.cell(1, 2 + col) 
                            cell.value = self.table.horizontalHeaderItem(col).text()

                            self.style_header_horizontal(cell, self.table.horizontalHeaderItem(col))

                            worksheet.column_dimensions[openpyxl.utils.get_column_letter(2 + col)].width = self.table.horizontalHeader().sectionSize(col) / dpi_horizontal * 13 + 3

                        worksheet.column_dimensions[openpyxl.utils.get_column_letter(1)].width = self.table.verticalHeader().width() / dpi_horizontal * 13 + 3

                        # Vertical Headers
                        for row_cell, row_item in enumerate(self.rows_export):
                            cell = worksheet.cell(2 + row_cell, 1)
                            cell.value = self.table.verticalHeaderItem(row_item).text()

                            self.style_header_vertical(cell, self.table.verticalHeaderItem(row_item))

                        # Cells
                        for row_cell, row_item in enumerate(self.rows_export):
                            for col in range(self.table.columnCount()):
                                cell = worksheet.cell(2 + row_cell, 2 + col)

                                cell_val = self.table.item(row_item, col).text()
                                # Remove illegal characters
                                cell_val = re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', cell_val)
                                cell.value = cell_val

                                if (col in self.table.headers_int or
                                    col in self.table.headers_float or
                                    col in self.table.headers_pct):
                                    self.style_cell_num(cell, self.table.item(row_item, col))
                                else:
                                    self.style_cell_text(cell, self.table.item(row_item, col))

                            self.progress_updated.emit(self.tr(f'Exporting table ... ({row_cell + 1} / {len_rows})'))

                    # Row Height
                    worksheet.row_dimensions[1].height = self.table.horizontalHeader().height() / dpi_vertical * 72

                    for i, _ in enumerate(worksheet.rows):
                        worksheet.row_dimensions[2 + i].height = self.table.verticalHeader().sectionSize(0) / dpi_vertical * 72

                    self.progress_updated.emit(self.tr(f'Saving file ...'))

                    workbook.save(self.file_path)
            elif self.file_type == self.tr('Word Document (*.docx)'):
                # Concordancer
                if self.table.tab == 'concordancer':
                    outputs = []

                    doc = docx.Document()

                    for i, row in enumerate(self.rows_export):
                        output = []

                        # Zapping
                        if settings_concordancer['zapping']:
                            # Discard position information
                            if settings_concordancer['discard_position_info']:
                                for j, col in enumerate(range(3)):
                                    if self.table.item(row, col):
                                        cell_text = self.table.item(row, col).text()
                                    else:
                                        cell_text = self.table.cellWidget(row, col).text()
                                        cell_text = wl_text_utils.html_to_text(cell_text)

                                    output.append(cell_text)
                            else:
                                if self.table.item(row, col):
                                    cell_text = self.table.item(row, col).text()
                                else:
                                    cell_text = self.table.cellWidget(row, col).text()
                                    cell_text = wl_text_utils.html_to_text(cell_text)

                                output.append(cell_text)

                            output[1] = settings_concordancer['placeholder'] * settings_concordancer['replace_keywords_with']

                            if settings_concordancer['add_line_nums']:
                                output.insert(0, f'{i + 1}. ')

                            outputs.append(output)
                        else:
                            for j, col in enumerate(range(3)):
                                cell_text = self.table.cellWidget(row, col).text()
                                cell_text = wl_text_utils.html_to_text(cell_text)

                                output.append(cell_text)

                        if not settings_concordancer['zapping']:
                            para = doc.add_paragraph(' '.join(output))
                            para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

                            self.progress_updated.emit(self.tr(f'Exporting table ... ({i + 1} / {len_rows})'))

                    # Randomize outputs
                    if settings_concordancer['zapping'] and settings_concordancer['randomize_outputs']:
                        random.shuffle(outputs)

                        # Re-order line numbers
                        if settings_concordancer['add_line_nums']:
                            for i, _ in enumerate(outputs):
                                outputs[i][0] = f'{i + 1}. '

                        for i, para in enumerate(outputs):
                            para = doc.add_paragraph(' '.join(para))
                            para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

                        self.progress_updated.emit(self.tr(f'Exporting table ... ({i + 1} / {len_rows})'))

                    self.progress_updated.emit(self.tr(f'Saving file ...'))

                    doc.save(self.file_path)
                # Concordancer (Parallel Mode)
                elif self.table.tab == 'concordancer_parallel':
                    # Source file
                    doc = docx.Document()

                    for i, row in enumerate(self.rows_export):
                        output = []

                        for j, col in enumerate(range(3)):
                            cell_text = self.table.linked_tables[0].cellWidget(row, col).text()
                            cell_text = wl_text_utils.html_to_text(cell_text)

                            output.append(cell_text)

                        para = doc.add_paragraph(' '.join(output))
                        para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

                        self.progress_updated.emit(self.tr(f'Exporting table ... ({i + 1} / {len_rows * 2})'))

                    doc.save(file_path_src)

                    # Target file
                    doc = docx.Document()

                    for i, row in enumerate(self.rows_export):
                        output = self.table.cellWidget(row, 0).text()

                        para = doc.add_paragraph(output)
                        para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

                        self.progress_updated.emit(self.tr(f'Exporting table ... ({len_rows + i + 1} / {len_rows * 2})'))

                    self.progress_updated.emit(self.tr(f'Saving file ...'))

                    doc.save(file_path_tgt)
                
            self.main.settings_custom['export']['tables']['default_path'] = wl_misc.get_normalized_dir(self.file_path)
            self.main.settings_custom['export']['tables']['default_type'] = self.file_type

            export_success = True
        except PermissionError:
            export_success = False

        self.worker_done.emit(export_success, self.file_path)

    def remove_illegal_chars(self, text):
        return re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', text)

    def style_header_horizontal(self, cell, item):
        cell.font = openpyxl.styles.Font(
            name = item.font().family(),
            size = 8,
            bold = True,
            color = 'FFFFFF'
        )

        if self.table.header_orientation == 'horizontal':
            cell.fill = openpyxl.styles.PatternFill(
                fill_type = 'solid',
                fgColor = '5C88C5'
            )
        else:
            cell.fill = openpyxl.styles.PatternFill(
                fill_type = 'solid',
                fgColor = '888888'
            )

        cell.alignment = openpyxl.styles.Alignment(
            horizontal = 'center',
            vertical = 'center',
            wrap_text = True
        )

    def style_header_vertical(self, cell, item):
        cell.font = openpyxl.styles.Font(
            name = item.font().family(),
            size = 8,
            bold = True,
            color = 'FFFFFF'
        )

        if self.table.header_orientation == 'horizontal':
            cell.fill = openpyxl.styles.PatternFill(
                fill_type = 'solid',
                fgColor = '888888'
            )

            cell.alignment = openpyxl.styles.Alignment(
                horizontal = 'right',
                vertical = 'center',
                wrap_text = True
          )
        else:
            cell.fill = openpyxl.styles.PatternFill(
                fill_type = 'solid',
                fgColor = '5C88C5'
            )

            cell.alignment = openpyxl.styles.Alignment(
                horizontal = 'left',
                vertical = 'center',
                wrap_text = True
            )

    def style_cell_text(self, cell, item):
        cell.font = openpyxl.styles.Font(
            name = item.font().family(),
            size = 8,
            color = '292929'
        )

        cell.alignment = openpyxl.styles.Alignment(
            horizontal = 'left',
            vertical = 'center',
            wrap_text = True
        )

    def style_cell_num(self, cell, item):
        cell.font = openpyxl.styles.Font(
            name = item.font().family(),
            size = 8,
            color = '292929'
        )

        cell.alignment = openpyxl.styles.Alignment(
            horizontal = 'right',
            vertical = 'center',
            wrap_text = True
        )

class Wl_Table_Item(QTableWidgetItem):
    def read_data(self):
        if (self.column() in self.tableWidget().headers_int or
            self.column() in self.tableWidget().headers_float or
            self.column() in self.tableWidget().headers_pct):
            return self.val
        else:
            return self.text()

    def __lt__(self, other):
        return self.read_data() < other.read_data()

class Wl_Table_Item_Error(QTableWidgetItem):
    def read_data(self):
        return self.text()

    def __lt__(self, other):
        return self.read_data() < other.read_data()

class Wl_Table(QTableWidget):
    def __init__(
        self, parent,
        headers, header_orientation = 'horizontal',
        cols_stretch = None,
        drag_drop_enabled = False
    ):
        self.main = wl_misc.find_wl_main(parent)

        self.headers = headers
        self.header_orientation = header_orientation
        self.cols_stretch = cols_stretch or []

        self.settings = self.main.settings_custom

        if header_orientation == 'horizontal':
            super().__init__(1, len(self.headers), parent)

            self.setHorizontalHeaderLabels(self.headers)
        else:
            super().__init__(len(self.headers), 1, parent)

            self.setVerticalHeaderLabels(self.headers)

        self.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)

        for col in self.find_col(self.cols_stretch):
            self.horizontalHeader().setSectionResizeMode(col, QHeaderView.Stretch)

        self.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.setSelectionBehavior(QAbstractItemView.SelectRows)

        if drag_drop_enabled:
            self.setDragEnabled(True)
            self.setAcceptDrops(True)
            self.viewport().setAcceptDrops(True)
            self.setDragDropMode(QAbstractItemView.InternalMove)
            self.setDragDropOverwriteMode(False)

        self.horizontalHeader().setHighlightSections(False)
        self.verticalHeader().setHighlightSections(False)

        self.setVerticalScrollMode(QAbstractItemView.ScrollPerPixel)
        self.setHorizontalScrollMode(QAbstractItemView.ScrollPerPixel)

        if self.header_orientation == 'horizontal':
            self.setStyleSheet(''' 
                QTableView {
                    outline: none;
                    color: #292929;
                }

                QTableView::item:hover {
                    background-color: #EEE;
                    color: #292929;
                }
                QTableView::item:selected {
                    background-color: #EEE;
                    color: #292929;
                }

                QHeaderView::section {
                    color: #FFF;
                    font-weight: bold;
                }

                QHeaderView::section:horizontal {
                    background-color: #5C88C5;
                }
                QHeaderView::section:horizontal:hover {
                    background-color: #3265B2;
                }
                QHeaderView::section:horizontal:pressed {
                    background-color: #3265B2;
                }

                QHeaderView::section:vertical {
                    background-color: #737373;
                }
                QHeaderView::section:vertical:hover {
                    background-color: #606060;
                }
                QHeaderView::section:vertical:pressed {
                    background-color: #606060;
                }
            ''')
        else:
            self.setStyleSheet('''
                QTableView {
                    outline: none;
                    color: #292929;
                }

                QTableView::item:hover {
                    background-color: #EEE;
                }
                QTableView::item:selected {
                    background-color: #EEE;
                    color: #292929;
                }

                QHeaderView::section {
                    color: #FFF;
                    font-weight: bold;
                }

                QHeaderView::section:horizontal {
                    background-color: #888;
                }
                QHeaderView::section:horizontal:hover {
                    background-color: #777;
                }
                QHeaderView::section:horizontal:pressed {
                    background-color: #666;
                }

                QHeaderView::section:vertical {
                    background-color: #5C88C5;
                }
                QHeaderView::section:vertical:hover {
                    background-color: #3265B2;
                }
                QHeaderView::section:vertical:pressed {
                    background-color: #264E8C;
                }
            ''')

        self.itemChanged.connect(self.item_changed)

    def dropEvent(self, event):
        if self.indexAt(event.pos()).row() == -1:
            row_dropped = self.rowCount()
        else:
            row_dropped = self.indexAt(event.pos()).row()

        selected_rows = self.get_selected_rows()

        self.blockSignals(True)

        rows_dragged = []

        for row in selected_rows:
            rows_dragged.append([])

            for col in range(self.columnCount()):
                if self.cellWidget(row, col):
                    rows_dragged[-1].append(self.cellWidget(row, col))
                else:
                    rows_dragged[-1].append(self.takeItem(row, col))

        for row in reversed(selected_rows):
            self.removeRow(row)

            if row < row_dropped:
                row_dropped -= 1

        for row, items in enumerate(rows_dragged):
            self.insertRow(row_dropped + row)

            for col, item in enumerate(items):
                if isinstance(item, QTableWidgetItem):
                    self.setItem(row_dropped + row, col, item)

                    self.item(row_dropped + row, col).setSelected(True)
                elif isinstance(item, QComboBox):
                    item_combo_box = wl_box.Wl_Combo_Box(self)
                    item_combo_box.addItems([item.itemText(i) for i in range(item.count())])
                    item_combo_box.setCurrentText(item.currentText())

                    self.setCellWidget(row_dropped + row, col, item_combo_box)
                elif isinstance(item, QLineEdit):
                    item_line_edit = QLineEdit(self)
                    item_line_edit.setText(item.text())

                    self.setCellWidget(row_dropped + row, col, item_line_edit)

        self.blockSignals(False)

        self.itemChanged.emit(self.item(0, 0))

        event.accept()

    def item_changed(self):
        cols_stretch = self.find_col(self.cols_stretch)

        self.resizeRowsToContents()

        for i in range(self.columnCount()):
            if i not in cols_stretch:
                self.resizeColumnToContents(i)

    def append_rows(self, labels):
        len_labels = len(labels)

        self.setRowCount(self.rowCount() + len_labels)

        for i, label in zip(range(self.rowCount() - len_labels, self.rowCount()), labels):
            self.setVerticalHeaderItem(i, QTableWidgetItem(label))

    def insert_col(self, i, label):
        super().insertColumn(i)

        self.setHorizontalHeaderItem(i, QTableWidgetItem(label))

    def export_selected(self):
        rows_export = sorted({index.row() for index in self.selectedIndexes()})

        self.export_all(rows_export = rows_export)

    def export_all(self, rows_export = None):
        def update_gui(export_success, file_path):
            self.results_exported = True

            if export_success:
                wl_msg_box.wl_msg_box_export_table_success(self.main, file_path)
            else:
                wl_msg_box.wl_msg_box_export_table_error(self.main, file_path)

        default_dir = self.main.settings_custom['export']['tables']['default_path']

        # Work Area
        if 'tab' in self.__dict__:
            if self.main.settings_custom['work_area_cur'] == 'Concordancer':
                if self.main.settings_custom['concordancer']['zapping_settings']['zapping']:
                    (file_path,
                     file_type) = QFileDialog.getSaveFileName(
                        self,
                        self.tr('Export Table'),
                        os.path.join(wl_checking_misc.check_dir(default_dir), 'Wordless_results_' + self.tab),
                        self.tr('Word Document (*.docx)'),
                        self.main.settings_custom['export']['tables']['default_type']
                    )
                else:
                    (file_path,
                     file_type) = QFileDialog.getSaveFileName(
                        self,
                        self.tr('Export Table'),
                        os.path.join(wl_checking_misc.check_dir(default_dir), 'Wordless_results_' + self.tab),
                        ';;'.join(self.main.settings_global['file_types']['export_tables_concordancer']),
                        self.main.settings_custom['export']['tables']['default_type']
                    )
            else:
                (file_path,
                 file_type) = QFileDialog.getSaveFileName(
                    self,
                    self.tr('Export Table'),
                    os.path.join(wl_checking_misc.check_dir(default_dir), 'Wordless_results_' + self.tab),
                    ';;'.join(self.main.settings_global['file_types']['export_tables']),
                    self.main.settings_custom['export']['tables']['default_type']
                )
        # Search terms, stop word lists, etc.
        else:
            (file_path,
             file_type) = QFileDialog.getSaveFileName(
                self,
                self.tr('Export Table'),
                os.path.join(wl_checking_misc.check_dir(default_dir), 'Wordless_import_error'),
                ';;'.join(self.main.settings_global['file_types']['export_tables']),
                self.main.settings_custom['export']['tables']['default_type']
            )

        if file_path:
            dialog_progress = wl_dialog_misc.Wl_Dialog_Progress_Export_Table(self.main)

            worker_export_table = Wl_Worker_Export_Table(
                self.main,
                dialog_progress = dialog_progress,
                update_gui = update_gui,
                table = self,
                file_path = file_path,
                file_type = file_type,
                rows_export = rows_export or []
            )

            thread_export_table = wl_threading.Wl_Thread(worker_export_table)
            thread_export_table.start_worker()

    def clear_table(self, header_count = 1):
        self.clearContents()

        if self.header_orientation == 'horizontal':
            self.setColumnCount(len(self.headers))
            self.setRowCount(header_count)

            self.setHorizontalHeaderLabels(self.headers)
        else:
            self.setRowCount(len(self.headers))
            self.setColumnCount(header_count)

            self.setVerticalHeaderLabels(self.headers)

        self.item_changed()

    def get_selected_rows(self):
        return sorted(set([index.row() for index in self.selectedIndexes()]))

    def find_row(self, text):
        def find(text):
            for row in range(self.rowCount()):
                if self.verticalHeaderItem(row).text() == text:
                    return row

        if type(text) == list:
            return [find(text_item) for text_item in text]
        else:
            return find(text)

    def find_rows(self, text):
        return [
            row
            for row in range(self.columnCount())
            if text in self.verticalHeaderItem(row).text()
        ]

    def find_col(self, text):
        def find(text):
            for col in range(self.columnCount()):
                if self.horizontalHeaderItem(col).text() == text:
                    return col

        if type(text) == list:
            return [find(text_item) for text_item in text]
        else:
            return find(text)

    def find_cols(self, text):
        return [
            col
            for col in range(self.columnCount())
            if text in self.horizontalHeaderItem(col).text()
        ]

    def find_header(self, text):
        if self.header_orientation == 'horizontal':
            return self.find_col(text)
        else:
            return self.find_row(text)

    def find_headers(self, text):
        if self.header_orientation == 'horizontal':
            return self.find_cols(text)
        else:
            return self.find_rows(text)

class Wl_Table_Error(Wl_Table):
    def __init__(self, main, headers):
        super().__init__(main, headers)

        self.name = 'error'

class Wl_Table_Data(Wl_Table):
    def __init__(
        self, main, tab,
        headers, header_orientation = 'horizontal',
        headers_int = None, headers_float = None,
        headers_pct = None, headers_cumulative = None, cols_breakdown = None,
        cols_stretch = None,
        sorting_enabled = False,
        linked_tables = None
    ):
        super().__init__(
            main, headers, header_orientation,
            cols_stretch,
            drag_drop_enabled = False
        )

        self.tab = tab

        self.headers_int_old = headers_int or []
        self.headers_float_old = headers_float or []
        self.headers_pct_old = headers_pct or []
        self.headers_cumulative_old = headers_cumulative or []
        self.cols_breakdown_old = cols_breakdown or []

        self.sorting_enabled = sorting_enabled
        self.linked_tables = linked_tables or []

        if sorting_enabled:
            self.setSortingEnabled(True)

            if header_orientation == 'horizontal':
                self.horizontalHeader().sortIndicatorChanged.connect(self.sorting_changed)
            else:
                self.verticalHeader().sortIndicatorChanged.connect(self.sorting_changed)

        self.itemChanged.connect(self.item_changed)
        self.itemSelectionChanged.connect(self.selection_changed)

        self.button_export_selected = QPushButton(self.tr('Export Selected...'), self)
        self.button_export_all = QPushButton(self.tr('Export All...'), self)
        self.button_clear = QPushButton(self.tr('Clear'), self)

        self.button_export_selected.clicked.connect(self.export_selected)
        self.button_export_all.clicked.connect(self.export_all)
        self.button_clear.clicked.connect(lambda: self.clear_table(confirm = True))

        self.clear_table()

    def item_changed(self):
        rows_visible = len([i for i in range(self.rowCount()) if not self.isRowHidden(i)])

        if [i for i in range(self.columnCount()) if self.item(0, i)] and rows_visible:
            self.button_export_all.setEnabled(True)
        else:
            self.button_export_all.setEnabled(False)

        super().item_changed()

        self.selection_changed()

    def selection_changed(self):
        for table in [self] + self.linked_tables:
            if table.selectedIndexes() and [i for i in range(table.columnCount()) if table.item(0, i)]:
                self.button_export_selected.setEnabled(True)

                break
            else:
                self.button_export_selected.setEnabled(False)

    def sorting_changed(self, logicalIndex, order):
        if [i for i in range(self.columnCount()) if self.item(0, i)]:
            self.update_ranks()

            if self.show_cumulative:
                self.toggle_cumulative()

    def append_rows(self, labels):
        headers_int = [self.verticalHeaderItem(row).text() for row in self.headers_int]
        headers_float = [self.verticalHeaderItem(row).text() for row in self.headers_float]
        headers_pct = [self.verticalHeaderItem(row).text() for row in self.headers_pct]
        headers_cumulative = [self.verticalHeaderItem(row).text() for row in self.headers_cumulative]

        super().append_rows([label for label, _, _, _, _ in labels])
        
        for label, is_int, is_float, is_pct, is_cumulative in labels:
            if is_int:
                headers_int.append(label)
            if is_float:
                headers_float.append(label)
            if is_pct:
                headers_pct.append(label)
            if is_cumulative:
                headers_cumulative.append(label)

        self.headers_int = set(self.find_row(headers_int))
        self.headers_float = set(self.find_row(headers_float))
        self.headers_pct = set(self.find_row(headers_pct))
        self.headers_cumulative = set(self.find_row(headers_cumulative))

    def insert_col(
        self, i, label,
        is_int = False, is_float = False,
        is_pct = False, is_cumulative = False, is_breakdown = False
    ):
        if self.header_orientation == 'horizontal':
            headers_int = [self.horizontalHeaderItem(col).text() for col in self.headers_int]
            headers_float = [self.horizontalHeaderItem(col).text() for col in self.headers_float]
            headers_pct = [self.horizontalHeaderItem(col).text() for col in self.headers_pct]
            headers_cumulative = [self.horizontalHeaderItem(col).text() for col in self.headers_cumulative]

        cols_breakdown = [self.horizontalHeaderItem(col).text() for col in self.cols_breakdown]

        super().insert_col(i, label)

        if is_int:
            headers_int += [label]
        if is_float:
            headers_float += [label]
        if is_pct:
            headers_pct += [label]
        if is_cumulative:
            headers_cumulative += [label]
        if is_breakdown:
            cols_breakdown += [label]

        if self.header_orientation == 'horizontal':
            self.headers_int = set(self.find_col(headers_int))
            self.headers_float = set(self.find_col(headers_float))
            self.headers_pct = set(self.find_col(headers_pct))
            self.headers_cumulative = set(self.find_col(headers_cumulative))
        
        self.cols_breakdown = set(self.find_col(cols_breakdown))

    def set_item_num(self, row, col, val, total = -1):
        if self.header_orientation == 'horizontal':
            header = col
        else:
            header = row

        # Integers
        if header in self.headers_int:
            val = int(val)

            item = Wl_Table_Item(str(val))
        # Floats
        elif header in self.headers_float:
            val = float(val)
            precision = self.main.settings_custom['data']['precision_decimal']

            item = Wl_Table_Item(f'{val:.{precision}f}')
        # Percentages
        elif header in self.headers_pct:
            if total > 0:
                val = val / total
            # Handle zero division error
            elif total == 0:
                val = 0
            # Set values directly
            elif total == -1:
                val = float(val)

            precision = self.main.settings_custom['data']['precision_pct']

            item = Wl_Table_Item(f'{val:.{precision}%}')

        item.val = val

        item.setFont(QFont('Consolas'))
        item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)

        super().setItem(row, col, item)

    def set_item_num_val(self, row, col, val):
        if self.header_orientation == 'horizontal':
            header = col
        else:
            header = row

        item = self.item(row, col)

        # Integers
        if header in self.headers_int:
            item.setText(str(val))
        # Floats
        elif header in self.headers_float:
            val = float(val)
            precision = self.main.settings_custom['data']['precision_decimal']

            item.setText(f'{val:.{precision}f}')
        # Percentages
        elif header in self.headers_pct:
            precision = self.main.settings_custom['data']['precision_pct']

            item.setText(f'{val:.{precision}%}')

        item.val = val

    def set_item_error(self, row, col, text):
        item = Wl_Table_Item_Error(text)

        item_font = QFont('Consolas')
        item_font.setBold(True)

        item.setFont(item_font)
        item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)

        super().setItem(row, col, item)

    def update_ranks(self):
        data_prev = ''
        rank_prev = 1
        rank_next = 1

        sort_section = self.horizontalHeader().sortIndicatorSection()
        sort_order = self.horizontalHeader().sortIndicatorOrder()

        col_rank = self.find_col(self.tr('Rank'))

        self.sortItems(sort_section, sort_order)

        if sort_section != col_rank:
            self.blockSignals(True)
            self.setSortingEnabled(False)
            self.setUpdatesEnabled(False)

            for row in range(self.rowCount()):
                if not self.isRowHidden(row):
                    data_cur = self.item(row, sort_section).read_data()

                    if self.main.settings_custom['data']['continue_numbering_after_ties']:
                        if data_cur == data_prev:
                            self.item(row, col_rank).val = rank_prev
                            self.item(row, col_rank).setText(str(rank_prev))
                        else:
                            self.item(row, col_rank).val = rank_next
                            self.item(row, col_rank).setText(str(rank_next))

                            rank_prev = rank_next
                            rank_next += 1
                            
                        data_prev = data_cur
                    else:
                        if data_cur == data_prev:
                            self.item(row, col_rank).val = rank_prev
                            self.item(row, col_rank).setText(str(rank_prev))
                        else:
                            self.item(row, col_rank).val = rank_next
                            self.item(row, col_rank).setText(str(rank_next))

                            rank_prev = rank_next

                        rank_next += 1
                        data_prev = data_cur

            self.blockSignals(False)
            self.setSortingEnabled(True)
            self.setUpdatesEnabled(True)

    def toggle_pct(self):
        self.setUpdatesEnabled(False)

        if self.header_orientation == 'horizontal':
            if self.show_pct:
                for col in self.headers_pct:
                    self.showColumn(col)
            else:
                for col in self.headers_pct:
                    self.hideColumn(col)
        else:
            if self.show_pct:
                for row in self.headers_pct:
                    self.showRow(row)
            else:
                for row in self.headers_pct:
                    self.hideRow(row)

        self.setUpdatesEnabled(True)

    def toggle_cumulative(self):
        precision_decimal = self.main.settings_custom['data']['precision_decimal']
        precision_pct = self.main.settings_custom['data']['precision_pct']

        # Boost performance
        self.sortItems(self.horizontalHeader().sortIndicatorSection(), self.horizontalHeader().sortIndicatorOrder())

        self.blockSignals(True)
        self.setSortingEnabled(False)
        self.setUpdatesEnabled(False)

        if self.header_orientation == 'horizontal':
            if self.show_cumulative:
                for col in self.headers_cumulative:
                    val_cumulative = 0

                    # Integers
                    if col in self.headers_int:
                        for row in range(self.rowCount()):
                            if not self.isRowHidden(row):
                                item = self.item(row, col)

                                val_cumulative += item.val
                                item.setText(str(val_cumulative))
                    # Floats
                    elif col in self.headers_float:
                        for row in range(self.rowCount()):
                            if not self.isRowHidden(row):
                                item = self.item(row, col)

                                val_cumulative += item.val
                                item.setText(f'{val_cumulative:.{precision_decimal}}')
                    # Percentages
                    elif col in self.headers_pct:
                        for row in range(self.rowCount()):
                            if not self.isRowHidden(row):
                                item = self.item(row, col)

                                val_cumulative += item.val
                                item.setText(f'{val_cumulative:.{precision_pct}%}')
            else:
                for col in self.headers_cumulative:
                    # Integers
                    if col in self.headers_int:
                        for row in range(self.rowCount()):
                            if not self.isRowHidden(row):
                                item = self.item(row, col)

                                item.setText(str(item.val))
                    # Floats
                    elif col in self.headers_float:
                        for row in range(self.rowCount()):
                            if not self.isRowHidden(row):
                                item = self.item(row, col)

                                item.setText(f'{item.val:.{precision_decimal}}')
                    # Percentages
                    elif col in self.headers_pct:
                        for row in range(self.rowCount()):
                            if not self.isRowHidden(row):
                                item = self.item(row, col)

                                item.setText(f'{item.val:.{precision_pct}%}')
        else:
            if self.show_cumulative:
                for row in self.headers_cumulative:
                    val_cumulative = 0

                    for col in range(self.columnCount() - 1):
                        item = self.item(row, col)

                        if not self.isColumnHidden(col) and not isinstance(item, Wl_Table_Item_Error):
                            val_cumulative += item.val

                            # Integers
                            if row in self.headers_int:
                                item.setText(str(val_cumulative))
                            # Floats
                            elif row in self.headers_float:
                                item.setText(f'{val_cumulative:.{precision_decimal}}')
                            # Percentages
                            elif row in self.headers_pct:
                                item.setText(f'{val_cumulative:.{precision_pct}%}')
            else:
                for row in self.headers_cumulative:
                    for col in range(self.columnCount() - 1):
                        item = self.item(row, col)

                        if not self.isColumnHidden(col) and not isinstance(item, Wl_Table_Item_Error):
                            # Integers
                            if row in self.headers_int:
                                item.setText(str(item.val))
                            # Floats
                            elif row in self.headers_float:
                                item.setText(f'{item.val:.{precision_decimal}}')
                            # Percentages
                            elif row in self.headers_pct:
                                item.setText(f'{item.val:.{precision_pct}%}')

        self.blockSignals(False)
        self.setUpdatesEnabled(True)

        if self.sorting_enabled:
            self.setSortingEnabled(True)

    def toggle_breakdown(self):
        self.setUpdatesEnabled(False)

        if self.show_breakdown:
            for col in self.cols_breakdown:
                self.showColumn(col)
        else:
            for col in self.cols_breakdown:
                self.hideColumn(col)

        self.setUpdatesEnabled(True)

    def filter_table(self):
        rank_prev = 1
        rank_next = 1
        data_prev = ''

        self.hide()
        self.setUpdatesEnabled(False)
        
        for i, row_filter in enumerate(self.row_filters):
            if row_filter:
                self.showRow(i)
            else:
                self.hideRow(i)

        self.setUpdatesEnabled(True)
        self.show()

        self.toggle_cumulative()
        self.update_ranks()

        self.itemChanged.emit(self.item(0, 0))

    def clear_table(self, count_headers = 1, confirm = False):
        confirmed = True

        # Ask for confirmation if results have not been exported
        if confirm:
            if not self.results_exported and self.item(0, 0) or self.cellWidget(0, 0):
                dialog_clear_table = wl_dialog_misc.WL_Dialog_Clear_Table(self.main)
                result = dialog_clear_table.exec_()

                if result == QDialog.Rejected:
                    confirmed = False

        if confirmed:
            for table in [self] + self.linked_tables:
                table.clearContents()

                if table.header_orientation == 'horizontal':
                    table.horizontalHeader().blockSignals(True)

                    table.setColumnCount(len(table.headers))
                    table.setRowCount(count_headers)

                    table.setHorizontalHeaderLabels(table.headers)

                    table.horizontalHeader().blockSignals(False)

                    table.horizontalHeader().sectionCountChanged.emit(0, count_headers)
                else:
                    table.verticalHeader().blockSignals(True)

                    table.setRowCount(len(table.headers))
                    table.setColumnCount(count_headers)

                    table.setVerticalHeaderLabels(table.headers)

                    table.verticalHeader().blockSignals(False)

                    table.verticalHeader().sectionCountChanged.emit(0, count_headers)

                for i in range(table.rowCount()):
                    table.showRow(i)

                for i in range(table.columnCount()):
                    table.showColumn(i)

                table.headers_int = set(table.find_header(table.headers_int_old))
                table.headers_float = set(table.find_header(table.headers_float_old))
                table.headers_pct = set(table.find_header(table.headers_pct_old))
                table.headers_cumulative = set(table.find_header(table.headers_cumulative_old))
                table.cols_breakdown = set(table.find_col(table.cols_breakdown_old))

                table.results_exported = False

                table.itemChanged.emit(table.item(0, 0))

        return confirmed

    def add_linked_table(self, table):
        self.linked_tables.append(table)

        table.itemSelectionChanged.connect(self.selection_changed)

    def add_linked_tables(self, tables):
        for table in tables:
            self.add_linked_table(table)

class Wl_Table_Data_Search(Wl_Table_Data):
    def __init__(
        self, main, tab,
        headers, header_orientation = 'horizontal',
        headers_int = None, headers_float = None,
        headers_pct = None, headers_cumulative = None, cols_breakdown = None,
        cols_stretch = None,
        sorting_enabled = False
    ):
        super().__init__(
            main, tab,
            headers, header_orientation,
            headers_int, headers_float,
            headers_pct, headers_cumulative, cols_breakdown,
            cols_stretch,
            sorting_enabled
        )

        self.label_number_results = QLabel()
        self.button_results_search = wl_button.Wl_Button_Results_Search(
            self,
            tab = self.tab,
            table = self
        )

        self.itemChanged.connect(self.results_changed)

        self.results_changed()

    def results_changed(self):
        rows_visible = len([i for i in range(self.rowCount()) if not self.isRowHidden(i)])

        if [i for i in range(self.columnCount()) if self.item(0, i)] and rows_visible:
            self.label_number_results.setText(self.tr(f'Number of Results: {rows_visible}'))

            self.button_results_search.setEnabled(True)
        else:
            self.label_number_results.setText(self.tr('Number of Results: 0'))

            self.button_results_search.setEnabled(False)

class Wl_Table_Data_Sort_Search(Wl_Table_Data):
    def __init__(
        self, main, tab,
        headers, header_orientation = 'horizontal',
        headers_int = None, headers_float = None,
        headers_pct = None, headers_cumulative = None, cols_breakdown = None,
        cols_stretch = None,
        sorting_enabled = False
    ):
        super().__init__(
            main, tab,
            headers, header_orientation,
            headers_int, headers_float,
            headers_pct, headers_cumulative, cols_breakdown,
            cols_stretch,
            sorting_enabled
        )

        self.label_number_results = QLabel()
        self.button_results_sort = wl_button.Wl_Button_Results_Sort(
            self,
            table = self
        )
        self.button_results_search = wl_button.Wl_Button_Results_Search(
            self,
            tab = self.tab,
            table = self
        )

        self.itemChanged.connect(self.results_changed)

        self.results_changed()

    def results_changed(self):
        rows_visible = len([i for i in range(self.rowCount()) if not self.isRowHidden(i)])

        if [i for i in range(self.columnCount()) if self.item(0, i)] and rows_visible:
            self.label_number_results.setText(self.tr(f'Number of Results: {rows_visible}'))

            self.button_results_sort.setEnabled(True)
            self.button_results_search.setEnabled(True)
        else:
            self.label_number_results.setText(self.tr('Number of Results: 0'))

            self.button_results_sort.setEnabled(False)
            self.button_results_search.setEnabled(False)

    def add_tables(self, tables):
        self.button_results_sort.add_tables(tables)
        self.button_results_search.add_tables(tables)

class Wl_Table_Data_Filter_Search(Wl_Table_Data):
    def __init__(
        self, main, tab,
        headers, header_orientation = 'horizontal',
        headers_int = None, headers_float = None,
        headers_pct = None, headers_cumulative = None, cols_breakdown = None,
        cols_stretch = None,
        sorting_enabled = False
    ):
        super().__init__(
            main, tab,
            headers, header_orientation,
            headers_int, headers_float,
            headers_pct, headers_cumulative, cols_breakdown,
            cols_stretch,
            sorting_enabled
        )

        self.label_number_results = QLabel()
        self.button_results_filter = wl_button.Wl_Button_Results_Filter(
            self,
            tab = self.tab,
            table = self
        )
        self.button_results_search = wl_button.Wl_Button_Results_Search(
            self,
            tab = self.tab,
            table = self
        )

        self.itemChanged.connect(self.results_changed)

        self.results_changed()

    def results_changed(self):
        rows_visible = len([i for i in range(self.rowCount()) if not self.isRowHidden(i)])

        if [i for i in range(self.columnCount()) if self.item(0, i)] and rows_visible:
            self.label_number_results.setText(self.tr(f'Number of Results: {rows_visible}'))

            self.button_results_filter.setEnabled(True)
            self.button_results_search.setEnabled(True)
        else:
            self.label_number_results.setText(self.tr('Number of Results: 0'))

            self.button_results_filter.setEnabled(False)
            self.button_results_search.setEnabled(False)

class Wl_Table_Tags(Wl_Table):
    def __init__(self, main):
        super().__init__(
            main,
            headers = [
                main.tr('Type'),
                main.tr('Level'),
                main.tr('Opening Tag'),
                main.tr('Closing Tag'),
                main.tr('Preview')
            ],
            header_orientation = 'horizontal',
            drag_drop_enabled = True
        )

        self.verticalHeader().setHidden(True)
        self.setFixedHeight(125)

        self.itemChanged.connect(self.item_changed)
        self.itemSelectionChanged.connect(self.selection_changed)

        self.button_add = QPushButton(self.tr('Add'), self)
        self.button_insert = QPushButton(self.tr('Insert'), self)
        self.button_remove = QPushButton(self.tr('Remove'), self)
        self.button_clear = QPushButton(self.tr('Clear'), self)
        self.button_reset = QPushButton(self.tr('Reset'), self)

        self.button_add.clicked.connect(self.add_item)
        self.button_insert.clicked.connect(self.insert_item)
        self.button_remove.clicked.connect(self.remove_item)
        self.button_clear.clicked.connect(lambda: self.clear_table(0))
        self.button_reset.clicked.connect(self.reset_table)

        self.reset_table()

    def item_changed(self, item = None):
        self.blockSignals(True)

        for row in range(self.rowCount()):
            opening_tag_widget = self.cellWidget(row, 2)

            # Opening Tag
            if re.search(r'^\s*$', opening_tag_widget.text()):
                wl_msg_box.wl_msg_box_empty_opening_tag(self.main)

                opening_tag_widget.setText(opening_tag_widget.text_old)
                opening_tag_widget.setFocus()

                self.blockSignals(False)

                return

        # Check for duplicate tags
        if item:
            item_row = item.row()
            item_col = item.column()

            if item_col == 2:
                item_widget = self.cellWidget(item_row, 2)

                for row in range(self.rowCount()):
                    if row != item_row:
                        if self.cellWidget(row, 2).text() == item_widget.text():
                            wl_msg_box.wl_msg_box_duplicate_tags(self.main)

                            item_widget.setText(item_widget.text_old)
                            item_widget.setFocus()

                            self.blockSignals(False)

                            return

        for row in range(self.rowCount()):
            type_text = self.cellWidget(row, 0).currentText()
            opening_tag_text = self.cellWidget(row, 2).text()
            closing_tag = self.item(row, 3)
            preview = self.item(row, 4)

            # Closing Tag
            if type_text == self.tr('Embedded'):
                closing_tag.setText(self.tr('N/A'))
            elif type_text == self.tr('Non-embedded'):
                # Add a "/" before the first non-punctuation character
                re_non_punc = re.search(r'\w|\*', opening_tag_text)

                if re_non_punc:
                    i_non_punc = re_non_punc.start()
                else:
                    i_non_punc = 1

                closing_tag.setText(f'{opening_tag_text[:i_non_punc]}/{opening_tag_text[i_non_punc:]}')

            # Preview
            if type_text == self.tr('Embedded'):
                preview.setText(self.tr(f'token{opening_tag_text}TAG'))
            elif type_text == self.tr('Non-embedded'):
                preview.setText(self.tr(f'{opening_tag_text}token{self.item(row, 3).text()}'))

        self.blockSignals(False)

        if self.rowCount():
            self.button_remove.setEnabled(True)
            self.button_clear.setEnabled(True)
        else:
            self.button_remove.setEnabled(False)
            self.button_clear.setEnabled(False)

        self.selection_changed()

    def selection_changed(self):
        if self.selectedIndexes():
            self.button_insert.setEnabled(True)

            if self.rowCount():
                self.button_remove.setEnabled(True)
            else:
                self.button_remove.setEnabled(False)
        else:
            self.button_insert.setEnabled(False)
            self.button_remove.setEnabled(False)

    def _new_item_type(self):
        new_item_type = wl_box.Wl_Combo_Box(self)

        new_item_type.addItems([
            self.tr('Embedded'),
            self.tr('Non-embedded')
        ])

        return new_item_type

    def _new_item_level(self):
        pass

    def _new_item_opening_tag(self, item_row):
        i = 1

        # Check for duplicate tag names if there are more than 1 row after the new row is added to the table
        duplicate = True

        while duplicate:
            if self.rowCount() > 1:
                for j in range(self.rowCount()):
                    if j != item_row and self.cellWidget(j, 2).text() == f'Tag_{i}':
                        i += 1

                        break
                    elif j == self.rowCount() - 1:
                        duplicate = False
            else:
                duplicate = False

        new_item = QLineEdit(f'Tag_{i}')

        new_item.text_old = new_item.text()
        
        return new_item

    def add_item(self, texts = []):
        self.blockSignals(True)

        self.setRowCount(self.rowCount() + 1)

        row = self.rowCount() - 1
        
        self.setCellWidget(row, 0, self._new_item_type())
        self.setCellWidget(row, 1, self._new_item_level())
        self.setCellWidget(row, 2, self._new_item_opening_tag(row))
        self.setItem(row, 0, QTableWidgetItem())
        self.setItem(row, 1, QTableWidgetItem())
        self.setItem(row, 2, QTableWidgetItem())

        self.setItem(row, 3, QTableWidgetItem())
        self.setItem(row, 4, QTableWidgetItem())

        if texts:
            self.cellWidget(row, 0).setCurrentText(texts[0])
            self.cellWidget(row, 1).setCurrentText(texts[1])
            self.cellWidget(row, 2).setText(texts[2])

        self.cellWidget(row, 0).currentTextChanged.connect(lambda: self.item_changed(item = self.item(row, 0)))
        self.cellWidget(row, 1).currentTextChanged.connect(lambda: self.item_changed(item = self.item(row, 1)))
        self.cellWidget(row, 2).editingFinished.connect(lambda: self.item_changed(item = self.item(row, 2)))

        self.blockSignals(False)

        self.item_changed()

    def insert_item(self):
        self.blockSignals(True)

        row = self.selectedIndexes()[0].row()

        self.insertRow(row)

        self.setCellWidget(row, 0, self._new_item_type())
        self.setCellWidget(row, 1, self._new_item_level())
        self.setCellWidget(row, 2, self._new_item_opening_tag(row))
        self.setItem(row, 0, QTableWidgetItem())
        self.setItem(row, 1, QTableWidgetItem())
        self.setItem(row, 2, QTableWidgetItem())

        self.setItem(row, 3, QTableWidgetItem())
        self.setItem(row, 4, QTableWidgetItem())

        self.cellWidget(row, 0).currentTextChanged.connect(lambda: self.item_changed(item = self.item(row, 0)))
        self.cellWidget(row, 1).currentTextChanged.connect(lambda: self.item_changed(item = self.item(row, 1)))
        self.cellWidget(row, 2).editingFinished.connect(lambda: self.item_changed(item = self.item(row, 2)))

        self.blockSignals(False)

        self.item_changed()

    def remove_item(self):
        self.blockSignals(True)

        for i in reversed(self.get_selected_rows()):
            self.removeRow(i)

        self.blockSignals(False)

        self.item_changed()

    def reset_table(self):
        pass

    def get_tags(self):
        tags = []

        for row in range(self.rowCount()):
            tags.append([
                self.cellWidget(row, 0).currentText(),
                self.cellWidget(row, 1).currentText(),
                self.cellWidget(row, 2).text(),
                self.item(row, 3).text()
            ])

        return tags
