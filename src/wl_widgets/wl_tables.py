# ----------------------------------------------------------------------
# Wordless: Widgets - Tables
# Copyright (C) 2018-2022  Ye Lei (叶磊)
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <http://www.gnu.org/licenses/>.
# ----------------------------------------------------------------------

import csv
import os
import random
import re

import docx
import openpyxl
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *

from wl_checking import wl_checking_misc
from wl_dialogs import wl_dialogs_misc, wl_msg_boxes
from wl_nlp import wl_nlp_utils
from wl_utils import wl_misc, wl_threading
from wl_widgets import wl_buttons

class Wl_Worker_Exp_Table(wl_threading.Wl_Worker):
    worker_done = pyqtSignal(bool, str)

    def run(self):
        if 'headers_int' not in self.table.__dict__:
            self.table.headers_int = []
        if 'headers_float' not in self.table.__dict__:
            self.table.headers_float = []
        if 'headers_pct' not in self.table.__dict__:
            self.table.headers_pct = []

        settings_concordancer = self.main.settings_custom['concordancer']['zapping_settings']

        # Check file permissions
        try:
            if not self.rows_to_exp:
                self.rows_to_exp = list(range(self.table.model().rowCount()))

            file_path_src = re.sub(r'\.([a-z]+?)$', r'_source.\1', self.file_path)
            file_path_tgt = re.sub(r'\.([a-z]+?)$', r'_target.\1', self.file_path)

            len_rows = len(self.rows_to_exp)

            # CSV files
            if self.file_type == self.tr('CSV File (*.csv)'):
                encoding = self.main.settings_custom['exp']['tables']['default_encoding']

                # Concordancer
                if self.table.tab == 'concordancer':
                    with open(self.file_path, 'w', encoding = encoding, newline = '') as f:
                        csv_writer = csv.writer(f)

                        # Horizontal Headers
                        csv_writer.writerow([
                            self.table.model().horizontalHeaderItem(col).text().strip()
                            for col in range(self.table.model().columnCount())
                        ])

                        # Cells
                        for i, row in enumerate(self.rows_to_exp):
                            row_to_exp = []

                            for col in range(self.table.model().columnCount()):
                                if self.table.model().item(row, col):
                                    cell_text = self.table.model().item(row, col).text()
                                else:
                                    cell_text = self.table.indexWidget(self.table.model().index(row, col)).text()
                                    cell_text = wl_nlp_utils.html_to_text(cell_text)

                                row_to_exp.append(cell_text)

                            csv_writer.writerow(row_to_exp)

                            self.progress_updated.emit(self.tr(f'Exporting table... ({i + 1} / {len_rows})'))
                # Concordancer (Parallel Mode)
                elif self.table.tab == 'concordancer_parallel':
                    # Source file
                    with open(file_path_src, 'w', encoding = encoding, newline = '') as f:
                        csv_writer = csv.writer(f)

                        # Horizontal Headers
                        csv_writer.writerow([
                            self.table.linked_tables[0].model().horizontalHeaderItem(col).text().strip()
                            for col in range(self.table.linked_tables[0].model().columnCount())
                        ])

                        # Cells
                        for i, row in enumerate(self.rows_to_exp):
                            row_to_exp = []

                            for col in range(self.table.linked_tables[0].model().columnCount()):
                                if self.table.linked_tables[0].model().item(row, col):
                                    cell_text = self.table.linked_tables[0].model().item(row, col).text()
                                else:
                                    cell_text = self.table.linked_tables[0].indexWidget(self.table.linked_tables[0].model().index(row, col)).text()
                                    cell_text = wl_nlp_utils.html_to_text(cell_text)

                                row_to_exp.append(cell_text)

                            csv_writer.writerow(row_to_exp)

                            self.progress_updated.emit(self.tr(f'Exporting table... ({i + 1} / {len_rows * 2})'))

                    # Target file
                    with open(file_path_tgt, 'w', encoding = encoding, newline = '') as f:
                        csv_writer = csv.writer(f)

                        # Horizontal Headers
                        csv_writer.writerow([
                            self.table.model().horizontalHeaderItem(col).text().strip()
                            for col in range(self.table.model().columnCount())
                        ])

                        # Cells
                        for i, row in enumerate(self.rows_to_exp):
                            row_to_exp = []

                            for col in range(self.table.model().columnCount()):
                                if self.table.model().item(row, col):
                                    cell_text = self.table.model().item(row, col).text()
                                else:
                                    cell_text = self.table.indexWidget(self.table.model().index(row, col)).text()
                                    cell_text = wl_nlp_utils.html_to_text(cell_text)

                                row_to_exp.append(cell_text)

                            csv_writer.writerow(row_to_exp)

                            self.progress_updated.emit(self.tr(f'Exporting table... ({len_rows + i + 1} / {len_rows * 2})'))
                else:
                    with open(self.file_path, 'w', encoding = encoding, newline = '') as f:
                        csv_writer = csv.writer(f)

                        if self.table.header_orientation == 'hor':
                            # Horizontal Headers
                            csv_writer.writerow([
                                self.table.model().horizontalHeaderItem(col).text().strip()
                                for col in range(self.table.model().columnCount())
                            ])

                            # Cells
                            for i, row in enumerate(self.rows_to_exp):
                                row_to_exp = []

                                for col in range(self.table.model().columnCount()):
                                    row_to_exp.append(self.table.model().item(row, col).text().strip())

                                csv_writer.writerow(row_to_exp)

                                self.progress_updated.emit(self.tr(f'Exporting table ... ({i + 1} / {len_rows})'))
                        else:
                            # Horizontal Headers
                            csv_writer.writerow(
                                ['']
                                + [
                                    self.table.model().horizontalHeaderItem(col).text().strip()
                                    for col in range(self.table.model().columnCount())
                                ]
                            )

                            # Vertical Headers & Cells
                            for i, row in enumerate(self.rows_to_exp):
                                row_to_exp = [self.table.model().verticalHeaderItem(row).text().strip()]

                                for col in range(self.table.model().columnCount()):
                                    row_to_exp.append(self.table.model().item(row, col).text().strip())

                                csv_writer.writerow(row_to_exp)

                                self.progress_updated.emit(self.tr(f'Exporting table ... ({i + 1} / {len_rows})'))
            # Excel workbooks
            elif self.file_type == self.tr('Excel Workbook (*.xlsx)'):
                dpi_horizontal = QApplication.primaryScreen().logicalDotsPerInchX()
                dpi_vertical = QApplication.primaryScreen().logicalDotsPerInchY()

                # Concordancer
                if self.table.tab == 'concordancer':
                    workbook = openpyxl.Workbook()
                    worksheet = workbook.active

                    worksheet.freeze_panes = 'A2'

                    dpi_horizontal = QApplication.primaryScreen().logicalDotsPerInchX()
                    dpi_vertical = QApplication.primaryScreen().logicalDotsPerInchY()

                    # Horizontal Headers
                    for col in range(self.table.model().columnCount()):
                        cell = worksheet.cell(1, 1 + col)
                        cell.value = self.table.model().horizontalHeaderItem(col).text()

                        self.style_header_horizontal(cell, self.table.model().horizontalHeaderItem(col))

                        worksheet.column_dimensions[openpyxl.utils.get_column_letter(1 + col)].width = self.table.horizontalHeader().sectionSize(col) / dpi_horizontal * 13 + 3

                    # Cells
                    for row_cell, row_item in enumerate(self.rows_to_exp):
                        for col in range(self.table.model().columnCount()):
                            # Left
                            if col == 0:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_nlp_utils.html_to_text(self.table.indexWidget(self.table.model().index(row_item, col)).text())
                                cell_val = self.remove_illegal_chars(cell_val)
                                cell.value = cell_val

                                self.style_cell_concordancer_left(cell, self.table.indexWidget(self.table.model().index(row_item, col)))
                            # Node
                            elif col == 1:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_nlp_utils.html_to_text(self.table.indexWidget(self.table.model().index(row_item, col)).text())
                                cell_val = self.remove_illegal_chars(cell_val)
                                cell.value = cell_val

                                self.style_cell_concordancer_node(cell, self.table.indexWidget(self.table.model().index(row_item, col)))
                            # Right
                            elif col == 2:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_nlp_utils.html_to_text(self.table.indexWidget(self.table.model().index(row_item, col)).text())
                                cell_val = self.remove_illegal_chars(cell_val)
                                cell.value = cell_val

                                self.style_cell_concordancer_right(cell, self.table.indexWidget(self.table.model().index(row_item, col)))
                            else:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = self.table.model().item(row_item, col).text()
                                cell_val = self.remove_illegal_chars(cell_val)
                                cell.value = cell_val

                                if (
                                    col in self.table.headers_int
                                    or col in self.table.headers_float
                                    or col in self.table.headers_pct
                                ):
                                    self.style_cell_num(cell, self.table.model().item(row_item, col))
                                else:
                                    self.style_cell_text(cell, self.table.model().item(row_item, col))

                            self.progress_updated.emit(self.tr(f'Exporting table... ({row_cell + 1} / {len_rows})'))

                    # Row Height
                    worksheet.row_dimensions[1].height = self.table.horizontalHeader().height() / dpi_vertical * 72

                    for i, _ in enumerate(worksheet.rows):
                        worksheet.row_dimensions[2 + i].height = self.table.verticalHeader().sectionSize(0) / dpi_vertical * 72

                    self.progress_updated.emit(self.tr('Saving file...'))

                    workbook.save(self.file_path)
                # Concordancer (Parallel Mode)
                elif self.table.tab == 'concordancer_parallel':
                    # Source file
                    workbook = openpyxl.Workbook()
                    worksheet = workbook.active

                    worksheet.freeze_panes = 'A2'

                    # Horizontal Headers
                    for col in range(self.table.linked_tables[0].model().columnCount()):
                        cell = worksheet.cell(1, 1 + col)
                        cell.value = self.table.linked_tables[0].model().horizontalHeaderItem(col).text()

                        self.style_header_horizontal(cell, self.table.linked_tables[0].model().horizontalHeaderItem(col))

                        worksheet.column_dimensions[openpyxl.utils.get_column_letter(1 + col)].width = self.table.linked_tables[0].horizontalHeader().sectionSize(col) / dpi_horizontal * 13 + 3

                    # Cells
                    for row_cell, row_item in enumerate(self.rows_to_exp):
                        for col in range(self.table.linked_tables[0].model().columnCount()):
                            # Left
                            if col == 0:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_nlp_utils.html_to_text(self.table.linked_tables[0].indexWidget(self.table.linked_tables[0].model().index(row_item, col)).text())
                                cell_val = self.remove_illegal_chars(cell_val)
                                cell.value = cell_val

                                self.style_cell_concordancer_left(cell, self.table.linked_tables[0].indexWidget(self.table.linked_tables[0].model().index(row_item, col)))
                            # Node
                            elif col == 1:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_nlp_utils.html_to_text(self.table.linked_tables[0].indexWidget(self.table.linked_tables[0].model().index(row_item, col)).text())
                                cell_val = self.remove_illegal_chars(cell_val)
                                cell.value = cell_val

                                self.style_cell_concordancer_node(cell, self.table.linked_tables[0].indexWidget(self.table.linked_tables[0].model().index(row_item, col)))
                            # Right
                            elif col == 2:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_nlp_utils.html_to_text(self.table.linked_tables[0].indexWidget(self.table.linked_tables[0].model().index(row_item, col)).text())
                                cell_val = self.remove_illegal_chars(cell_val)
                                cell.value = cell_val

                                self.style_cell_concordancer_right(cell, self.table.linked_tables[0].indexWidget(self.table.linked_tables[0].model().index(row_item, col)))
                            else:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = self.table.linked_tables[0].model().item(row_item, col).text()
                                cell_val = self.remove_illegal_chars(cell_val)
                                cell.value = cell_val

                                if (
                                    col in self.table.linked_tables[0].headers_int
                                    or col in self.table.linked_tables[0].headers_float
                                    or col in self.table.linked_tables[0].headers_pct
                                ):
                                    self.style_cell_num(cell, self.table.linked_tables[0].model().item(row_item, col))
                                else:
                                    self.style_cell_text(cell, self.table.linked_tables[0].model().item(row_item, col))

                            self.progress_updated.emit(self.tr(f'Exporting table... ({row_cell + 1} / {len_rows * 2})'))

                    # Row Height
                    worksheet.row_dimensions[1].height = self.table.linked_tables[0].horizontalHeader().height() / dpi_vertical * 72

                    for i, _ in enumerate(worksheet.rows):
                        worksheet.row_dimensions[2 + i].height = self.table.linked_tables[0].verticalHeader().sectionSize(0) / dpi_vertical * 72

                    self.progress_updated.emit(self.tr('Saving source file...'))

                    workbook.save(file_path_src)

                    # Source file
                    workbook = openpyxl.Workbook()
                    worksheet = workbook.active

                    worksheet.freeze_panes = 'A2'

                    # Horizontal Headers
                    for col in range(self.table.model().columnCount()):
                        cell = worksheet.cell(1, 1 + col)
                        cell.value = self.table.model().horizontalHeaderItem(col).text()

                        self.style_header_horizontal(cell, self.table.model().horizontalHeaderItem(col))

                        worksheet.column_dimensions[openpyxl.utils.get_column_letter(1 + col)].width = self.table.horizontalHeader().sectionSize(col) / dpi_horizontal * 13 + 3

                    # Cells
                    for row_cell, row_item in enumerate(self.rows_to_exp):
                        for col in range(self.table.model().columnCount()):
                            # Parallel Text
                            if col == 0:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = wl_nlp_utils.html_to_text(self.table.indexWidget(self.table.model().index(row_item, col)).text())
                                cell_val = self.remove_illegal_chars(cell_val)
                                cell.value = cell_val

                                self.style_cell_concordancer_parallel_node(cell, self.table.indexWidget(self.table.model().index(row_item, col)))
                            else:
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = self.table.model().item(row_item, col).text()
                                cell_val = self.remove_illegal_chars(cell_val)
                                cell.value = cell_val

                                if (
                                    col in self.table.headers_int
                                    or col in self.table.headers_float
                                    or col in self.table.headers_pct
                                ):
                                    self.style_cell_num(cell, self.table.model().item(row_item, col))
                                else:
                                    self.style_cell_text(cell, self.table.model().item(row_item, col))

                            self.progress_updated.emit(self.tr(f'Exporting table... ({len_rows + row_cell + 1} / {len_rows * 2})'))

                    # Row Height
                    worksheet.row_dimensions[1].height = self.table.horizontalHeader().height() / dpi_vertical * 72

                    for i, _ in enumerate(worksheet.rows):
                        worksheet.row_dimensions[2 + i].height = self.table.verticalHeader().sectionSize(0) / dpi_vertical * 72

                    self.progress_updated.emit(self.tr('Saving target file...'))

                    workbook.save(file_path_tgt)
                else:
                    workbook = openpyxl.Workbook()
                    worksheet = workbook.active

                    worksheet.freeze_panes = 'B2'

                    if self.table.header_orientation == 'hor':
                        # Horizontal Headers
                        for col in range(self.table.model().columnCount()):
                            cell = worksheet.cell(1, 1 + col)
                            cell.value = self.table.model().horizontalHeaderItem(col).text()

                            self.style_header_horizontal(cell, self.table.model().horizontalHeaderItem(col))

                            worksheet.column_dimensions[openpyxl.utils.get_column_letter(1 + col)].width = self.table.horizontalHeader().sectionSize(col) / dpi_horizontal * 13 + 3

                        # Cells
                        for row_cell, row_item in enumerate(self.rows_to_exp):
                            for col in range(self.table.model().columnCount()):
                                cell = worksheet.cell(2 + row_cell, 1 + col)

                                cell_val = self.table.model().item(row_item, col).text()
                                cell_val = self.remove_illegal_chars(cell_val)
                                cell.value = cell_val

                                if (
                                    col in self.table.headers_int
                                    or col in self.table.headers_float
                                    or col in self.table.headers_pct
                                ):
                                    self.style_cell_num(cell, self.table.model().item(row_item, col))
                                else:
                                    self.style_cell_text(cell, self.table.model().item(row_item, col))

                            self.progress_updated.emit(self.tr(f'Exporting table ... ({row_cell + 1} / {len_rows})'))
                    else:
                        # Horizontal Headers
                        for col in range(self.table.model().columnCount()):
                            cell = worksheet.cell(1, 2 + col)
                            cell.value = self.table.model().horizontalHeaderItem(col).text()

                            self.style_header_horizontal(cell, self.table.model().horizontalHeaderItem(col))

                            worksheet.column_dimensions[openpyxl.utils.get_column_letter(2 + col)].width = self.table.horizontalHeader().sectionSize(col) / dpi_horizontal * 13 + 3

                        worksheet.column_dimensions[openpyxl.utils.get_column_letter(1)].width = self.table.verticalHeader().width() / dpi_horizontal * 13 + 3

                        # Vertical Headers
                        for row_cell, row_item in enumerate(self.rows_to_exp):
                            cell = worksheet.cell(2 + row_cell, 1)
                            cell.value = self.table.model().verticalHeaderItem(row_item).text()

                            self.style_header_vertical(cell, self.table.model().verticalHeaderItem(row_item))

                        # Cells
                        for row_cell, row_item in enumerate(self.rows_to_exp):
                            for col in range(self.table.model().columnCount()):
                                cell = worksheet.cell(2 + row_cell, 2 + col)

                                cell_val = self.table.model().item(row_item, col).text()
                                cell_val = self.remove_illegal_chars(cell_val)
                                cell.value = cell_val

                                if (
                                    col in self.table.headers_int
                                    or col in self.table.headers_float
                                    or col in self.table.headers_pct
                                ):
                                    self.style_cell_num(cell, self.table.model().item(row_item, col))
                                else:
                                    self.style_cell_text(cell, self.table.model().item(row_item, col))

                            self.progress_updated.emit(self.tr(f'Exporting table ... ({row_cell + 1} / {len_rows})'))

                    # Row Height
                    worksheet.row_dimensions[1].height = self.table.horizontalHeader().height() / dpi_vertical * 72

                    for i, _ in enumerate(worksheet.rows):
                        worksheet.row_dimensions[2 + i].height = self.table.verticalHeader().sectionSize(0) / dpi_vertical * 72

                    self.progress_updated.emit(self.tr('Saving file ...'))

                    workbook.save(self.file_path)
            elif self.file_type == self.tr('Word Document (*.docx)'):
                # Concordancer
                if self.table.tab == 'concordancer':
                    outputs = []

                    doc = docx.Document()

                    for i, row in enumerate(self.rows_to_exp):
                        output = []

                        # Zapping
                        if settings_concordancer['zapping']:
                            # Discard position information
                            if settings_concordancer['discard_position_info']:
                                for j, col in enumerate(range(3)):
                                    if self.table.model().item(row, col):
                                        cell_text = self.table.model().item(row, col).text()
                                    else:
                                        cell_text = self.table.indexWidget(self.table.model().index(row, col)).text()
                                        cell_text = wl_nlp_utils.html_to_text(cell_text)

                                    output.append(cell_text)
                            else:
                                if self.table.item(row, col):
                                    cell_text = self.table.model().item(row, col).text()
                                else:
                                    cell_text = self.table.indexWidget(self.table.model().index(row, col)).text()
                                    cell_text = wl_nlp_utils.html_to_text(cell_text)

                                output.append(cell_text)

                            output[1] = settings_concordancer['placeholder'] * settings_concordancer['replace_keywords_with']

                            if settings_concordancer['add_line_nums']:
                                output.insert(0, f'{i + 1}. ')

                            outputs.append(output)
                        else:
                            for j, col in enumerate(range(3)):
                                cell_text = self.table.indexWidget(self.table.model().index(row, col)).text()
                                cell_text = wl_nlp_utils.html_to_text(cell_text)

                                output.append(cell_text)

                        if not settings_concordancer['zapping']:
                            para = doc.add_paragraph(' '.join(output))
                            para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

                            self.progress_updated.emit(self.tr(f'Exporting table ... ({i + 1} / {len_rows})'))

                    # Randomize outputs
                    if settings_concordancer['zapping'] and settings_concordancer['randomize_outputs']:
                        random.shuffle(outputs)

                        # Re-order line numbers
                        if settings_concordancer['add_line_nums']:
                            for i, _ in enumerate(outputs):
                                outputs[i][0] = f'{i + 1}. '

                        for i, para in enumerate(outputs):
                            para = doc.add_paragraph(' '.join(para))
                            para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

                        self.progress_updated.emit(self.tr(f'Exporting table... ({i + 1} / {len_rows})'))

                    self.progress_updated.emit(self.tr('Saving file...'))

                    doc.save(self.file_path)
                # Concordancer (Parallel Mode)
                elif self.table.tab == 'concordancer_parallel':
                    # Source file
                    doc = docx.Document()

                    for i, row in enumerate(self.rows_to_exp):
                        output = []

                        for j, col in enumerate(range(3)):
                            cell_text = self.table.linked_tables[0].indexWidget(self.table.linked_tables[0].model().index(row, col)).text()
                            cell_text = wl_nlp_utils.html_to_text(cell_text)

                            output.append(cell_text)

                        para = doc.add_paragraph(' '.join(output))
                        para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

                        self.progress_updated.emit(self.tr(f'Exporting table ... ({i + 1} / {len_rows * 2})'))

                    doc.save(file_path_src)

                    # Target file
                    doc = docx.Document()

                    for i, row in enumerate(self.rows_to_exp):
                        output = self.table.indexWidget(self.table.model().index(row, 0)).text()

                        para = doc.add_paragraph(output)
                        para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

                        self.progress_updated.emit(self.tr(f'Exporting table... ({len_rows + i + 1} / {len_rows * 2})'))

                    self.progress_updated.emit(self.tr('Saving file...'))

                    doc.save(file_path_tgt)

            self.main.settings_custom['exp']['tables']['default_path'] = wl_misc.get_normalized_dir(self.file_path)
            self.main.settings_custom['exp']['tables']['default_type'] = self.file_type

            exp_success = True
        except PermissionError:
            exp_success = False

        self.worker_done.emit(exp_success, self.file_path)

    # Remove illegal characters
    def remove_illegal_chars(self, text):
        return re.sub(openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE, '', text)

    def style_header_horizontal(self, cell, item):
        cell.font = openpyxl.styles.Font(
            name = item.font().family(),
            size = self.main.settings_custom['general']['font_settings']['font_size'] - 5,
            bold = True,
            color = 'FFFFFF'
        )

        if self.table.header_orientation == 'hor':
            cell.fill = openpyxl.styles.PatternFill(
                fill_type = 'solid',
                fgColor = '5C88C5'
            )
        else:
            cell.fill = openpyxl.styles.PatternFill(
                fill_type = 'solid',
                fgColor = '888888'
            )

        cell.alignment = openpyxl.styles.Alignment(
            horizontal = 'center',
            vertical = 'center',
            wrap_text = True
        )

    def style_header_vertical(self, cell, item):
        cell.font = openpyxl.styles.Font(
            name = item.font().family(),
            size = self.main.settings_custom['general']['font_settings']['font_size'] - 5,
            bold = True,
            color = 'FFFFFF'
        )

        if self.table.header_orientation == 'hor':
            cell.fill = openpyxl.styles.PatternFill(
                fill_type = 'solid',
                fgColor = '888888'
            )
            cell.alignment = openpyxl.styles.Alignment(
                horizontal = 'right',
                vertical = 'center',
                wrap_text = True
            )
        else:
            cell.fill = openpyxl.styles.PatternFill(
                fill_type = 'solid',
                fgColor = '5C88C5'
            )
            cell.alignment = openpyxl.styles.Alignment(
                horizontal = 'left',
                vertical = 'center',
                wrap_text = True
            )

    def style_cell_text(self, cell, item):
        cell.font = openpyxl.styles.Font(
            name = item.font().family(),
            size = self.main.settings_custom['general']['font_settings']['font_size'] - 5,
            color = '292929'
        )
        cell.alignment = openpyxl.styles.Alignment(
            horizontal = 'left',
            vertical = 'center',
            wrap_text = True
        )

    def style_cell_num(self, cell, item):
        cell.font = openpyxl.styles.Font(
            name = item.font().family(),
            size = self.main.settings_custom['general']['font_settings']['font_size'] - 5,
            color = '292929'
        )
        cell.alignment = openpyxl.styles.Alignment(
            horizontal = 'right',
            vertical = 'center',
            wrap_text = True
        )

    def style_cell_concordancer_node(self, cell, item):
        cell.font = openpyxl.styles.Font(
            name = item.font().family(),
            size = self.main.settings_custom['general']['font_settings']['font_size'] - 5,
            bold = True,
            color = 'FF0000'
        )
        cell.alignment = openpyxl.styles.Alignment(
            horizontal = 'center',
            vertical = 'center',
            wrap_text = True
        )

    def style_cell_concordancer_left(self, cell, item):
        cell.font = openpyxl.styles.Font(
            name = item.font().family(),
            size = self.main.settings_custom['general']['font_settings']['font_size'] - 5,
            color = '292929'
        )
        cell.alignment = openpyxl.styles.Alignment(
            horizontal = 'right',
            vertical = 'center',
            wrap_text = True
        )

    def style_cell_concordancer_right(self, cell, item):
        cell.font = openpyxl.styles.Font(
            name = item.font().family(),
            size = self.main.settings_custom['general']['font_settings']['font_size'] - 5,
            color = '292929'
        )
        cell.alignment = openpyxl.styles.Alignment(
            horizontal = 'left',
            vertical = 'center',
            wrap_text = True
        )

    def style_cell_concordancer_parallel_node(self, cell, item):
        cell.font = openpyxl.styles.Font(
            name = item.font().family(),
            size = self.main.settings_custom['general']['font_settings']['font_size'] - 5,
            color = '292929'
        )
        cell.alignment = openpyxl.styles.Alignment(
            horizontal = 'center',
            vertical = 'center',
            wrap_text = True
        )

class Wl_Item_Delegate_Uneditable(QStyledItemDelegate):
    def createEditor(self, parent, option, index):
        pass

class Wl_Table(QTableView):
    def __init__(
        self, parent,
        headers, header_orientation = 'hor',
        editable = False,
        drag_drop = False
    ):
        super().__init__(parent)

        self.main = wl_misc.find_wl_main(parent)

        self.headers = headers
        self.header_orientation = header_orientation

        self.settings = self.main.settings_custom

        model = QStandardItemModel()
        model.table = self

        self.setModel(model)

        if header_orientation == 'hor':
            self.model().setHorizontalHeaderLabels(self.headers)
        else:
            self.model().setVerticalHeaderLabels(self.headers)

        self.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)

        self.horizontalHeader().setHighlightSections(False)
        self.verticalHeader().setHighlightSections(False)

        self.setVerticalScrollMode(QAbstractItemView.ScrollPerPixel)
        self.setHorizontalScrollMode(QAbstractItemView.ScrollPerPixel)

        if editable:
            self.setEditTriggers(QAbstractItemView.DoubleClicked | QAbstractItemView.SelectedClicked)
        else:
            self.setEditTriggers(QAbstractItemView.NoEditTriggers)

        if drag_drop:
            self.setDragEnabled(True)
            self.setAcceptDrops(True)
            self.viewport().setAcceptDrops(True)
            self.setDragDropMode(QAbstractItemView.InternalMove)
            self.setDragDropOverwriteMode(False)

        self.setSelectionBehavior(QAbstractItemView.SelectRows)

        if self.header_orientation == 'hor':
            self.setStyleSheet('''
                QTableView {
                    outline: none;
                    color: #292929;
                }

                QTableView::item:hover {
                    background-color: #EEE;
                    color: #292929;
                }
                QTableView::item:selected {
                    background-color: #EEE;
                    color: #292929;
                }

                QHeaderView::section {
                    color: #FFF;
                    font-weight: bold;
                }

                QHeaderView::section:horizontal {
                    background-color: #5C88C5;
                }
                QHeaderView::section:horizontal:hover {
                    background-color: #3265B2;
                }
                QHeaderView::section:horizontal:pressed {
                    background-color: #3265B2;
                }

                QHeaderView::section:vertical {
                    background-color: #737373;
                }
                QHeaderView::section:vertical:hover {
                    background-color: #606060;
                }
                QHeaderView::section:vertical:pressed {
                    background-color: #606060;
                }
            ''')
        else:
            self.setStyleSheet('''
                QTableView {
                    outline: none;
                    color: #292929;
                }

                QTableView::item:hover {
                    background-color: #EEE;
                }
                QTableView::item:selected {
                    background-color: #EEE;
                    color: #292929;
                }

                QHeaderView::section {
                    color: #FFF;
                    font-weight: bold;
                }

                QHeaderView::section:horizontal {
                    background-color: #888;
                }
                QHeaderView::section:horizontal:hover {
                    background-color: #777;
                }
                QHeaderView::section:horizontal:pressed {
                    background-color: #666;
                }

                QHeaderView::section:vertical {
                    background-color: #5C88C5;
                }
                QHeaderView::section:vertical:hover {
                    background-color: #3265B2;
                }
                QHeaderView::section:vertical:pressed {
                    background-color: #264E8C;
                }
            ''')

        self.model().itemChanged.connect(self.item_changed)
        self.selectionModel().selectionChanged.connect(self.selection_changed)

    # There seems to be a bug with QAbstractItemView.InternalMove
    # See: https://bugreports.qt.io/browse/QTBUG-87057
    def dropEvent(self, event):
        if self.indexAt(event.pos()).row() == -1:
            row_dropped_on = self.model().rowCount()
        else:
            row_dropped_on = self.indexAt(event.pos()).row()

        data = []
        data_selected = []
        rows_selected = self.get_selected_rows()

        for row in range(self.model().rowCount()):
            data.append([])

            for col in range(self.model().columnCount()):
                item = self.model().takeItem(row, col)
                widget = self.indexWidget(self.model().index(row, col))

                if widget:
                    data[-1].append(type(widget)(widget.text(), self))
                else:
                    data[-1].append(item)

        for row in rows_selected:
            data_selected.append(data[row].copy())
            data[row] = []

        for row in reversed(data_selected):
            data.insert(row_dropped_on, row)

        data = [row for row in data if row]

        self.disable_updates()

        self.clr_table()
        self.model().setRowCount(len(data))
        self.model().setColumnCount(len(data[0]))

        for i, row in enumerate(data):
            for j, item in enumerate(row):
                if type(item) == QStandardItem:
                    self.model().setItem(i, j, item)
                else:
                    self.setIndexWidget(self.model().index(i, j), item)

        self.enable_updates()

        event.accept()

    def item_changed(self, item):
        if self.is_empty():
            self.setEnabled(False)
        else:
            self.setEnabled(True)

        self.resizeColumnsToContents()
        self.resizeRowsToContents()

        self.selectionModel().selectionChanged.emit(QItemSelection(), QItemSelection())

    def selection_changed(self, selected, deselected):
        pass

    def disable_updates(self):
        self.num_rows_old = self.model().rowCount()
        self.num_cols_old = self.model().columnCount()

        self.setUpdatesEnabled(False)
        self.blockSignals(True)
        self.horizontalHeader().blockSignals(True)
        self.verticalHeader().blockSignals(True)
        self.model().blockSignals(True)
        self.selectionModel().blockSignals(True)
        self.hide()

    def enable_updates(self, emit_signals = True):
        self.setUpdatesEnabled(True)
        self.blockSignals(False)
        self.horizontalHeader().blockSignals(False)
        self.verticalHeader().blockSignals(False)
        self.model().blockSignals(False)
        self.selectionModel().blockSignals(False)
        self.show()

        self.horizontalHeader().sectionCountChanged.emit(self.num_cols_old, self.model().columnCount())
        self.verticalHeader().sectionCountChanged.emit(self.num_rows_old, self.model().rowCount())

        if emit_signals:
            self.model().itemChanged.emit(QStandardItem())
            self.selectionModel().selectionChanged.emit(QItemSelection(), QItemSelection())

    def is_empty(self):
        empty = False

        if self.header_orientation == 'hor':
            if not any([
                    self.indexWidget(self.model().index(0, i)) or self.model().item(0, i)
                    for i in range(self.model().columnCount())
            ]):
                empty = True
        else:
            if not any([
                    self.indexWidget(self.model().index(i, 0)) or self.model().item(i, 0)
                    for i in range(self.model().rowCount())
            ]):
                empty = True

        return empty

    def get_header_labels_hor(self):
        return (self.model().headerData(row, Qt.Horizontal) for row in range(self.model().columnCount()))

    def get_header_labels_vert(self):
        return (self.model().headerData(col, Qt.Vertical) for col in range(self.model().rowCount()))

    def find_header_hor(self, text):
        return list(self.get_header_labels_hor()).index(text)

    def find_header_vert(self, text):
        return list(self.get_header_labels_vert()).index(text)

    def find_headers_hor(self, text):
        return [i for i, header in enumerate(self.get_header_labels_hor()) if text in header]

    def find_headers_vert(self, text):
        return [i for i, header in enumerate(self.get_header_labels_vert()) if text in header]

    def find_header(self, text):
        if self.header_orientation == 'hor':
            return self.find_header_hor(text = text)
        else:
            return self.find_header_vert(text = text)

    def find_headers(self, text):
        if self.header_orientation == 'hor':
            return self.find_headers_hor(text = text)
        else:
            return self.find_headers_vert(text = text)

    def add_header_hor(self, label):
        self.add_headers_hor([label])

    def add_header_vert(self, label):
        self.add_headers_vert([label])

    def add_headers_hor(self, labels):
        self.model().setHorizontalHeaderLabels(list(self.get_header_labels_hor()) + labels)

    def add_headers_vert(self, labels):
        self.model().setVerticalHeaderLabels(list(self.get_header_labels_vert()) + labels)

    def ins_header_hor(self, i, label):
        self.ins_headers_hor(i, [label])

    def ins_header_vert(self, i, label):
        self.ins_headers_vert(i, [label])

    def ins_headers_hor(self, i, labels):
        headers = list(self.get_header_labels_hor())
        headers[i:i] = labels

        self.model().setHorizontalHeaderLabels(headers)

    def ins_headers_vert(self, i, labels):
        headers = list(self.get_header_labels_vert())
        headers[i:i] = labels

        self.model().setVerticalHeaderLabels(headers)

    def get_selected_rows(self):
        return sorted({index.row() for index in self.selectionModel().selectedIndexes()})

    def get_selected_cols(self):
        return sorted({index.col() for index in self.selectionModel().selectedIndexes()})

    def _add_row(self, row = None, texts = None):
        if texts is None:
            texts = self.defaults_row

        if self.is_empty():
            self.clr_table(0)

        if row is None:
            self.model().appendRow([QStandardItem(text) for text in texts])
        else:
            self.model().insertRow(row, [QStandardItem(text) for text in texts])

        self.model().itemChanged.emit(QStandardItem())

    def add_row(self, texts = None):
        self._add_row(texts = texts)
        self.setCurrentIndex(self.model().index(self.model().rowCount() - 1, 0))

    def ins_row(self, texts = None):
        row = self.get_selected_rows()[0]

        self._add_row(row = row, texts = texts)
        self.setCurrentIndex(self.model().index(row, 0))

    def del_row(self):
        for i in reversed(self.get_selected_rows()):
            self.model().removeRow(i)

        self.model().itemChanged.emit(QStandardItem())

    def clr_table(self, num_headers = 1):
        self.model().clear()

        if self.header_orientation == 'hor':
            self.model().setHorizontalHeaderLabels(self.headers)
            self.model().setRowCount(num_headers)
        else:
            self.model().setVerticalHeaderLabels(self.headers)
            self.model().setColumnCount(num_headers)

        self.model().itemChanged.emit(QStandardItem())

    def exp_selected(self):
        self.exp_all(rows_to_exp = self.get_selected_rows())

    def exp_all(self, rows_to_exp = None):
        def update_gui(exp_success, file_path):
            self.results_saved = True

            if exp_success:
                wl_msg_boxes.Wl_Msg_Box_Info(
                    self.main,
                    title = self.tr('Export Completed'),
                    text = self.tr(f'''
                        <div>The table has been successfully exported to "{file_path}".</div>
                    ''')
                ).open()
            else:
                wl_msg_boxes.Wl_Msg_Box_Info(
                    self.main,
                    title = self.tr('Export Error'),
                    text = self.tr(f'''
                        <div>Access to "{file_path}" is denied, please specify another location or close the file and try again.</div>
                    ''')
                ).open()

        default_dir = self.main.settings_custom['exp']['tables']['default_path']

        # Search terms, stop word lists, file checking, etc.
        if self.tab == 'err':
            file_path, file_type = QFileDialog.getSaveFileName(
                self,
                self.tr('Export Table'),
                os.path.join(wl_checking_misc.check_dir(default_dir), 'wordless_error'),
                ';;'.join(self.main.settings_global['file_types']['exp_tables']),
                self.main.settings_custom['exp']['tables']['default_type']
            )
        # Work Area
        else:
            # Concordancer
            if self.tab in ['concordancer', 'concordancer_parallel']:
                if self.tab == 'concordancer' and self.main.settings_custom['concordancer']['zapping_settings']['zapping']:
                    file_path, file_type = QFileDialog.getSaveFileName(
                        self,
                        self.tr('Export Table'),
                        os.path.join(wl_checking_misc.check_dir(default_dir), f'wordless_results_{self.tab}'),
                        self.tr('Word Document (*.docx)')
                    )
                else:
                    file_path, file_type = QFileDialog.getSaveFileName(
                        self,
                        self.tr('Export Table'),
                        os.path.join(wl_checking_misc.check_dir(default_dir), f'wordless_results_{self.tab}'),
                        ';;'.join(self.main.settings_global['file_types']['exp_tables_concordancer']),
                        self.main.settings_custom['exp']['tables']['default_type']
                    )
            else:
                file_path, file_type = QFileDialog.getSaveFileName(
                    self,
                    self.tr('Export Table'),
                    os.path.join(wl_checking_misc.check_dir(default_dir), f'wordless_results_{self.tab}'),
                    ';;'.join(self.main.settings_global['file_types']['exp_tables']),
                    self.main.settings_custom['exp']['tables']['default_type']
                )

        if file_path:
            dialog_progress = wl_dialogs_misc.Wl_Dialog_Progress(self.main, text = self.tr('Exporting table...'))

            worker_exp_table = Wl_Worker_Exp_Table(
                self.main,
                dialog_progress = dialog_progress,
                update_gui = update_gui,
                table = self,
                file_path = file_path,
                file_type = file_type,
                rows_to_exp = rows_to_exp or []
            )

            thread_exp_table = wl_threading.Wl_Thread(worker_exp_table)
            thread_exp_table.start_worker()

class Wl_Table_Add_Ins_Del_Clr(Wl_Table):
    def __init__(self, parent, headers, defaults = None):
        super().__init__(
            parent = parent,
            headers = headers,
            editable = True,
            drag_drop = True
        )

        self.defaults = defaults or []

        self.button_add = QPushButton(self.tr('Add'), self)
        self.button_ins = QPushButton(self.tr('Insert'), self)
        self.button_del = QPushButton(self.tr('Remove'), self)
        self.button_clr = QPushButton(self.tr('Clear'), self)
        self.button_reset = QPushButton(self.tr('Reset'), self)

        self.button_add.clicked.connect(lambda: self.add_row())
        self.button_ins.clicked.connect(lambda: self.ins_row())
        self.button_del.clicked.connect(lambda: self.del_row())
        self.button_clr.clicked.connect(lambda: self.clr_table(0))
        self.button_reset.clicked.connect(lambda: self.reset_table())

    def item_changed(self, item):
        if self.model().rowCount():
            self.button_del.setEnabled(True)
            self.button_clr.setEnabled(True)
        else:
            self.button_del.setEnabled(False)
            self.button_clr.setEnabled(False)

        super().item_changed(item)

    def selection_changed(self, selected, deselected):
        if self.selectionModel().selectedIndexes():
            self.button_ins.setEnabled(True)

            if self.model().rowCount():
                self.button_del.setEnabled(True)
            else:
                self.button_del.setEnabled(False)
        else:
            self.button_ins.setEnabled(False)
            self.button_del.setEnabled(False)

    def reset_table(self):
        self.clr_table(0)

        for default_texts in self.defaults:
            self._add_row(texts = default_texts)

class Wl_Table_Item(QStandardItem):
    def read_data(self):
        if (
            self.column() in self.model().table.headers_int
            or self.column() in self.model().table.headers_float
            or self.column() in self.model().table.headers_pct
        ):
            return self.val
        else:
            return self.text()

    def __lt__(self, other):
        return self.read_data() < other.read_data()

class Wl_Table_Item_Error(QStandardItem):
    def read_data(self):
        return self.text()

    def __lt__(self, other):
        return self.read_data() < other.read_data()

class Wl_Table_Data(Wl_Table):
    def __init__(
        self, main, tab,
        headers, header_orientation = 'hor',
        headers_int = None, headers_float = None,
        headers_pct = None, headers_cumulative = None, cols_breakdown = None,
        sorting_enabled = False,
        linked_tables = None
    ):
        super().__init__(
            main, headers, header_orientation,
            editable = False,
            drag_drop = False
        )

        self.tab = tab

        self.headers_int_old = headers_int or {}
        self.headers_float_old = headers_float or {}
        self.headers_pct_old = headers_pct or {}
        self.headers_cumulative_old = headers_cumulative or {}
        self.cols_breakdown_old = cols_breakdown or {}

        self.sorting_enabled = sorting_enabled
        self.linked_tables = linked_tables or []

        if sorting_enabled:
            self.setSortingEnabled(True)

            if header_orientation == 'hor':
                self.horizontalHeader().sortIndicatorChanged.connect(self.sorting_changed)
            else:
                self.verticalHeader().sortIndicatorChanged.connect(self.sorting_changed)

        self.model().itemChanged.connect(self.item_changed)
        self.selectionModel().selectionChanged.connect(self.selection_changed)

        self.button_exp_selected = QPushButton(self.tr('Export Selected...'), self)
        self.button_exp_all = QPushButton(self.tr('Export All...'), self)
        self.button_clr = QPushButton(self.tr('Clear'), self)

        self.button_exp_selected.clicked.connect(self.exp_selected)
        self.button_exp_all.clicked.connect(self.exp_all)
        self.button_clr.clicked.connect(lambda: self.clr_table(confirm = True))

        self.clr_table()

    def item_changed(self, item):
        rows_visible = len([i for i in range(self.model().rowCount()) if not self.isRowHidden(i)])

        if [i for i in range(self.model().columnCount()) if self.model().item(0, i)] and rows_visible:
            self.button_exp_all.setEnabled(True)
        else:
            self.button_exp_all.setEnabled(False)

        if [i for i in range(self.model().columnCount()) if self.model().item(0, i)]:
            self.button_clr.setEnabled(True)
        else:
            self.button_clr.setEnabled(False)

        super().item_changed(item)

        self.selection_changed()

    def selection_changed(self):
        for table in [self] + self.linked_tables:
            if not self.is_empty() and [i for i in range(self.model().rowCount()) if not self.isRowHidden(i)] and table.selectionModel().selectedIndexes():
                self.button_exp_selected.setEnabled(True)

                break
            else:
                self.button_exp_selected.setEnabled(False)

    def sorting_changed(self, logicalIndex, order):
        if not self.is_empty():
            self.update_ranks()

            if self.show_cumulative:
                self.toggle_cumulative()

    def add_header_hor(
        self, label,
        is_int = False, is_float = False,
        is_pct = False, is_cumulative = False, is_breakdown = False
    ):
        self.add_headers_hor(
            [label],
            is_int = is_int, is_float = is_float,
            is_pct = is_pct, is_cumulative = is_cumulative, is_breakdown = is_breakdown
        )

    def add_header_vert(
        self, label,
        is_int = False, is_float = False,
        is_pct = False, is_cumulative = False
    ):
        self.add_headers_vert(
            [label],
            is_int = is_int, is_float = is_float,
            is_pct = is_pct, is_cumulative = is_cumulative
        )

    def add_headers_hor(
        self, labels,
        is_int = False, is_float = False,
        is_pct = False, is_cumulative = False, is_breakdown = False
    ):
        super().add_headers_hor(labels)

        cols_labels = {self.find_header_hor(label) for label in labels}

        if is_int:
            self.headers_int |= cols_labels
        if is_float:
            self.headers_float |= cols_labels
        if is_pct:
            self.headers_pct |= cols_labels
        if is_cumulative:
            self.headers_cumulative |= cols_labels
        if is_breakdown:
            self.cols_breakdown |= cols_labels

    def add_headers_vert(
        self, labels,
        is_int = False, is_float = False,
        is_pct = False, is_cumulative = False
    ):
        super().add_headers_vert(labels)

        rows_labels = {self.find_header_vert(label) for label in labels}

        if is_int:
            self.headers_int |= rows_labels
        if is_float:
            self.headers_float |= rows_labels
        if is_pct:
            self.headers_pct |= rows_labels
        if is_cumulative:
            self.headers_cumulative |= rows_labels

    def ins_header_hor(
        self, i, label,
        is_int = False, is_float = False,
        is_pct = False, is_cumulative = False, is_breakdown = False
    ):
        self.ins_headers_hor(
            i, [label],
            is_int = is_int, is_float = is_float,
            is_pct = is_pct, is_cumulative = is_cumulative, is_breakdown = is_breakdown
        )

    def ins_header_vert(
        self, i, label,
        is_int = False, is_float = False,
        is_pct = False, is_cumulative = False
    ):
        self.ins_headers_vert(
            i, [label],
            is_int = is_int, is_float = is_float,
            is_pct = is_pct, is_cumulative = is_cumulative
        )

    def ins_headers_hor(
        self, i, labels,
        is_int = False, is_float = False,
        is_pct = False, is_cumulative = False, is_breakdown = False
    ):
        # Re-calculate column indexes
        if self.header_orientation == 'hor':
            headers_int = [self.model().horizontalHeaderItem(col).text() for col in self.headers_int]
            headers_float = [self.model().horizontalHeaderItem(col).text() for col in self.headers_float]
            headers_pct = [self.model().horizontalHeaderItem(col).text() for col in self.headers_pct]
            headers_cumulative = [self.model().horizontalHeaderItem(col).text() for col in self.headers_cumulative]

        cols_breakdown = [self.model().horizontalHeaderItem(col).text() for col in self.cols_breakdown]

        super().ins_headers_hor(i, labels)

        if self.header_orientation == 'hor':
            if is_int:
                headers_int.extend(labels)
            if is_float:
                headers_float.extend(labels)
            if is_pct:
                headers_pct.extend(labels)
            if is_cumulative:
                headers_cumulative.extend(labels)

            self.headers_int = {self.find_header_hor(header) for header in headers_int}
            self.headers_float = {self.find_header_hor(header) for header in headers_float}
            self.headers_pct = {self.find_header_hor(header) for header in headers_pct}
            self.headers_cumulative = {self.find_header_hor(header) for header in headers_cumulative}

        if is_breakdown:
            cols_breakdown.extend(labels)

        self.cols_breakdown = {self.find_header_hor(header) for header in cols_breakdown}

    def ins_headers_vert(
        self, i, labels,
        is_int = False, is_float = False,
        is_pct = False, is_cumulative = False
    ):
        # Re-calculate row indexes
        headers_int = [self.model().verticalHeaderItem(row).text() for row in self.headers_int]
        headers_float = [self.model().verticalHeaderItem(row).text() for row in self.headers_float]
        headers_pct = [self.model().verticalHeaderItem(row).text() for row in self.headers_pct]
        headers_cumulative = [self.model().verticalHeaderItem(row).text() for row in self.headers_cumulative]

        super().ins_headers_vert(i, labels)

        if is_int:
            headers_int.extend(labels)
        if is_float:
            headers_float.extend(labels)
        if is_pct:
            headers_pct.extend(labels)
        if is_cumulative:
            headers_cumulative.extend(labels)

        self.headers_int = {self.find_header_vert(header) for header in headers_int}
        self.headers_float = {self.find_header_vert(header) for header in headers_float}
        self.headers_pct = {self.find_header_vert(header) for header in headers_pct}
        self.headers_cumulative = {self.find_header_vert(header) for header in headers_cumulative}

    def set_item_num(self, row, col, val, total = -1):
        if self.header_orientation == 'hor':
            header = col
        else:
            header = row

        # Integers
        if header in self.headers_int:
            val = int(val)

            item = Wl_Table_Item(str(val))
        # Floats
        elif header in self.headers_float:
            val = float(val)
            precision = self.main.settings_custom['data']['precision_decimal']

            item = Wl_Table_Item(f'{val:.{precision}f}')
        # Percentages
        elif header in self.headers_pct:
            if total > 0:
                val = val / total
            # Handle zero division error
            elif total == 0:
                val = 0
            # Set values directly
            elif total == -1:
                val = float(val)

            precision = self.main.settings_custom['data']['precision_pct']

            item = Wl_Table_Item(f'{val:.{precision}%}')

        item.val = val

        item.setFont(QFont('Consolas'))
        item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)

        self.model().setItem(row, col, item)

    def set_item_num_val(self, row, col, val):
        if self.header_orientation == 'hor':
            header = col
        else:
            header = row

        item = self.model().item(row, col)

        # Integers
        if header in self.headers_int:
            item.setText(str(val))
        # Floats
        elif header in self.headers_float:
            val = float(val)
            precision = self.main.settings_custom['data']['precision_decimal']

            item.setText(f'{val:.{precision}f}')
        # Percentages
        elif header in self.headers_pct:
            precision = self.main.settings_custom['data']['precision_pct']

            item.setText(f'{val:.{precision}%}')

        item.val = val

    def set_item_error(self, row, col, text):
        item = Wl_Table_Item_Error(text)

        item_font = QFont('Consolas')
        item_font.setBold(True)

        item.setFont(item_font)
        item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)

        self.model().setItem(row, col, item)

    def update_ranks(self):
        data_prev = ''
        rank_prev = 1
        rank_next = 1

        sort_section = self.horizontalHeader().sortIndicatorSection()
        sort_order = self.horizontalHeader().sortIndicatorOrder()

        col_rank = self.find_header_hor(self.tr('Rank'))

        self.sortByColumn(sort_section, sort_order)

        if sort_section != col_rank:
            self.disable_updates()

            for row in range(self.model().rowCount()):
                if not self.isRowHidden(row):
                    data_cur = self.model().item(row, sort_section).read_data()

                    if self.main.settings_custom['data']['continue_numbering_after_ties']:
                        if data_cur == data_prev:
                            self.model().item(row, col_rank).val = rank_prev
                            self.model().item(row, col_rank).setText(str(rank_prev))
                        else:
                            self.model().item(row, col_rank).val = rank_next
                            self.model().item(row, col_rank).setText(str(rank_next))

                            rank_prev = rank_next
                            rank_next += 1

                        data_prev = data_cur
                    else:
                        if data_cur == data_prev:
                            self.model().item(row, col_rank).val = rank_prev
                            self.model().item(row, col_rank).setText(str(rank_prev))
                        else:
                            self.model().item(row, col_rank).val = rank_next
                            self.model().item(row, col_rank).setText(str(rank_next))

                            rank_prev = rank_next

                        rank_next += 1
                        data_prev = data_cur

            self.enable_updates()

    def toggle_pct(self):
        self.disable_updates()

        if self.header_orientation == 'hor':
            if self.show_pct:
                for col in self.headers_pct:
                    self.showColumn(col)
            else:
                for col in self.headers_pct:
                    self.hideColumn(col)
        else:
            if self.show_pct:
                for row in self.headers_pct:
                    self.showRow(row)
            else:
                for row in self.headers_pct:
                    self.hideRow(row)

        self.enable_updates()

    def toggle_cumulative(self):
        precision_decimal = self.main.settings_custom['data']['precision_decimal']
        precision_pct = self.main.settings_custom['data']['precision_pct']

        # Boost performance
        self.sortByColumn(self.horizontalHeader().sortIndicatorSection(), self.horizontalHeader().sortIndicatorOrder())

        self.disable_updates()
        self.setSortingEnabled(False)

        if self.header_orientation == 'hor':
            if self.show_cumulative:
                for col in self.headers_cumulative:
                    val_cumulative = 0

                    # Integers
                    if col in self.headers_int:
                        for row in range(self.model().rowCount()):
                            if not self.isRowHidden(row):
                                item = self.model().item(row, col)

                                val_cumulative += item.val
                                item.setText(str(val_cumulative))
                    # Floats
                    elif col in self.headers_float:
                        for row in range(self.model().rowCount()):
                            if not self.isRowHidden(row):
                                item = self.model().item(row, col)

                                val_cumulative += item.val
                                item.setText(f'{val_cumulative:.{precision_decimal}}')
                    # Percentages
                    elif col in self.headers_pct:
                        for row in range(self.model().rowCount()):
                            if not self.isRowHidden(row):
                                item = self.model().item(row, col)

                                val_cumulative += item.val
                                item.setText(f'{val_cumulative:.{precision_pct}%}')
            else:
                for col in self.headers_cumulative:
                    # Integers
                    if col in self.headers_int:
                        for row in range(self.model().rowCount()):
                            if not self.isRowHidden(row):
                                item = self.model().item(row, col)

                                item.setText(str(item.val))
                    # Floats
                    elif col in self.headers_float:
                        for row in range(self.model().rowCount()):
                            if not self.isRowHidden(row):
                                item = self.model().item(row, col)

                                item.setText(f'{item.val:.{precision_decimal}}')
                    # Percentages
                    elif col in self.headers_pct:
                        for row in range(self.model().rowCount()):
                            if not self.isRowHidden(row):
                                item = self.model().item(row, col)

                                item.setText(f'{item.val:.{precision_pct}%}')
        else:
            if self.show_cumulative:
                for row in self.headers_cumulative:
                    val_cumulative = 0

                    for col in range(self.model().columnCount() - 1):
                        item = self.model().item(row, col)

                        if not self.isColumnHidden(col) and not isinstance(item, Wl_Table_Item_Error):
                            val_cumulative += item.val

                            # Integers
                            if row in self.headers_int:
                                item.setText(str(val_cumulative))
                            # Floats
                            elif row in self.headers_float:
                                item.setText(f'{val_cumulative:.{precision_decimal}}')
                            # Percentages
                            elif row in self.headers_pct:
                                item.setText(f'{val_cumulative:.{precision_pct}%}')
            else:
                for row in self.headers_cumulative:
                    for col in range(self.model().columnCount() - 1):
                        item = self.model().item(row, col)

                        if not self.isColumnHidden(col) and not isinstance(item, Wl_Table_Item_Error):
                            # Integers
                            if row in self.headers_int:
                                item.setText(str(item.val))
                            # Floats
                            elif row in self.headers_float:
                                item.setText(f'{item.val:.{precision_decimal}}')
                            # Percentages
                            elif row in self.headers_pct:
                                item.setText(f'{item.val:.{precision_pct}%}')

        self.enable_updates()

        if self.sorting_enabled:
            self.setSortingEnabled(True)

    def toggle_breakdown(self):
        self.disable_updates()

        if self.show_breakdown:
            for col in self.cols_breakdown:
                self.showColumn(col)
        else:
            for col in self.cols_breakdown:
                self.hideColumn(col)

        self.enable_updates()

    def filter_table(self):
        self.disable_updates()

        for i, row_filter in enumerate(self.row_filters):
            if row_filter:
                self.showRow(i)
            else:
                self.hideRow(i)

        self.enable_updates()

        self.toggle_cumulative()
        self.update_ranks()

    def clr_table(self, count_headers = 1, confirm = False):
        confirmed = True

        # Ask for confirmation if results have not been exported
        if confirm:
            if not self.is_empty() and not self.results_saved:
                dialog_clear_table = wl_dialogs_misc.WL_Dialog_Clear_Table(self.main)
                result = dialog_clear_table.exec_()

                if result == QDialog.Rejected:
                    confirmed = False

        if confirmed:
            for table in [self] + self.linked_tables:
                table.model().clear()

                if table.header_orientation == 'hor':
                    table.horizontalHeader().blockSignals(True)

                    table.model().setColumnCount(len(table.headers))
                    table.model().setRowCount(count_headers)

                    table.model().setHorizontalHeaderLabels(table.headers)

                    table.horizontalHeader().blockSignals(False)

                    table.horizontalHeader().sectionCountChanged.emit(0, count_headers)
                else:
                    table.verticalHeader().blockSignals(True)

                    table.model().setRowCount(len(table.headers))
                    table.model().setColumnCount(count_headers)

                    table.model().setVerticalHeaderLabels(table.headers)

                    table.verticalHeader().blockSignals(False)

                    table.verticalHeader().sectionCountChanged.emit(0, count_headers)

                for i in range(table.model().rowCount()):
                    table.showRow(i)

                for i in range(table.model().columnCount()):
                    table.showColumn(i)

                table.headers_int = {table.find_header(header) for header in table.headers_int_old}
                table.headers_float = {table.find_header(header) for header in table.headers_float_old}
                table.headers_pct = {table.find_header(header) for header in table.headers_pct_old}
                table.headers_cumulative = {table.find_header(header) for header in table.headers_cumulative_old}
                table.cols_breakdown = {table.find_header_hor(col) for col in table.cols_breakdown_old}

                table.results_saved = False

                table.model().itemChanged.emit(QStandardItem())

        return confirmed

    def add_linked_table(self, table):
        self.linked_tables.append(table)

        table.selectionModel().selectionChanged.connect(self.selection_changed)

    def add_linked_tables(self, tables):
        for table in tables:
            self.add_linked_table(table)

# Avoid circular imports
from wl_results import wl_results_filter, wl_results_search, wl_results_sort

class Wl_Table_Data_Search(Wl_Table_Data):
    def __init__(
        self, main, tab,
        headers, header_orientation = 'hor',
        headers_int = None, headers_float = None,
        headers_pct = None, headers_cumulative = None, cols_breakdown = None,
        sorting_enabled = False
    ):
        super().__init__(
            main, tab,
            headers, header_orientation,
            headers_int, headers_float,
            headers_pct, headers_cumulative, cols_breakdown,
            sorting_enabled
        )

        self.model().itemChanged.connect(self.results_changed)

        self.label_number_results = QLabel()
        self.button_results_search = wl_buttons.Wl_Button(
            self.tr('Search in Results'),
            self
        )
        self.dialog_results_search = wl_results_search.Wl_Dialog_Results_Search(
            self.main,
            tab = self.tab,
            table = self
        )

        self.button_results_search.setFixedWidth(150)

        self.button_results_search.clicked.connect(self.dialog_results_search.load)

        self.results_changed()

    def results_changed(self):
        rows_visible = len([i for i in range(self.rowCount()) if not self.isRowHidden(i)])

        if not self.is_empty() and rows_visible:
            self.label_number_results.setText(self.tr(f'Number of Results: {rows_visible}'))

            self.button_results_search.setEnabled(True)
        else:
            self.label_number_results.setText(self.tr('Number of Results: 0'))

            self.button_results_search.setEnabled(False)

class Wl_Table_Data_Sort_Search(Wl_Table_Data):
    def __init__(
        self, main, tab,
        headers, header_orientation = 'hor',
        headers_int = None, headers_float = None,
        headers_pct = None, headers_cumulative = None, cols_breakdown = None,
        sorting_enabled = False
    ):
        super().__init__(
            main, tab,
            headers, header_orientation,
            headers_int, headers_float,
            headers_pct, headers_cumulative, cols_breakdown,
            sorting_enabled
        )

        self.model().itemChanged.connect(self.results_changed)

        self.label_number_results = QLabel()
        self.button_results_search = wl_buttons.Wl_Button(
            self.tr('Search in Results'),
            self
        )
        self.button_results_sort = wl_buttons.Wl_Button(
            self.tr('Sort Results'),
            self
        )

        self.dialog_results_search = wl_results_search.Wl_Dialog_Results_Search(
            self.main,
            tab = self.tab,
            table = self
        )
        self.dialog_results_sort = wl_results_sort.Wl_Dialog_Results_Sort_Concordancer(
            self.main,
            table = self
        )

        self.button_results_search.setFixedWidth(150)
        self.button_results_sort.setFixedWidth(150)

        self.button_results_search.clicked.connect(self.dialog_results_search.load)
        self.button_results_sort.clicked.connect(self.dialog_results_sort.show)

        self.results_changed()

    def results_changed(self):
        rows_visible = len([i for i in range(self.model().rowCount()) if not self.isRowHidden(i)])

        if not self.is_empty() and rows_visible:
            self.label_number_results.setText(self.tr(f'Number of Results: {rows_visible}'))

            self.button_results_sort.setEnabled(True)
            self.button_results_search.setEnabled(True)
        else:
            self.label_number_results.setText(self.tr('Number of Results: 0'))

            self.button_results_sort.setEnabled(False)
            self.button_results_search.setEnabled(False)

    def add_tables(self, tables):
        self.dialog_results_sort.add_tables(tables)
        self.dialog_results_search.add_tables(tables)

class Wl_Table_Data_Filter_Search(Wl_Table_Data):
    def __init__(
        self, main, tab,
        headers, header_orientation = 'hor',
        headers_int = None, headers_float = None,
        headers_pct = None, headers_cumulative = None, cols_breakdown = None,
        sorting_enabled = False
    ):
        super().__init__(
            main, tab,
            headers, header_orientation,
            headers_int, headers_float,
            headers_pct, headers_cumulative, cols_breakdown,
            sorting_enabled
        )

        self.model().itemChanged.connect(self.results_changed)

        self.label_number_results = QLabel()
        self.button_results_filter = wl_buttons.Wl_Button(
            self.tr('Filter Results'),
            self
        )
        self.button_results_search = wl_buttons.Wl_Button(
            self.tr('Search in Results'),
            self
        )

        if self.tab in ['wordlist', 'ngram']:
            self.dialog_results_filter = wl_results_filter.Wl_Dialog_Results_Filter_Wordlist(self.main, tab = self.tab, table = self)
        elif self.tab in ['collocation', 'colligation']:
            self.dialog_results_filter = wl_results_filter.Wl_Dialog_Results_Filter_Collocation(self.main, tab = self.tab, table = self)
        elif self.tab == 'keyword':
            self.dialog_results_filter = wl_results_filter.Wl_Dialog_Results_Filter_Keyword(self.main, tab = self.tab, table = self)

        self.dialog_results_search = wl_results_search.Wl_Dialog_Results_Search(
            self.main,
            tab = self.tab,
            table = self
        )

        self.button_results_filter.setFixedWidth(150)
        self.button_results_search.setFixedWidth(150)

        self.button_results_filter.clicked.connect(self.dialog_results_filter.show)
        self.button_results_search.clicked.connect(self.dialog_results_search.load)

        self.results_changed()

    def results_changed(self):
        rows_visible = len([i for i in range(self.model().rowCount()) if not self.isRowHidden(i)])

        if not self.is_empty():
            self.label_number_results.setText(self.tr(f'Number of Results: {rows_visible}'))

            self.button_results_filter.setEnabled(True)
        else:
            self.label_number_results.setText(self.tr('Number of Results: 0'))

            self.button_results_filter.setEnabled(False)

        if not self.is_empty() and rows_visible:
            self.button_results_search.setEnabled(True)
        else:
            self.button_results_search.setEnabled(False)
