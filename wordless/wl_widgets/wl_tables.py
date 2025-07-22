# ----------------------------------------------------------------------
# Wordless: Widgets - Tables
# Copyright (C) 2018-2025  Ye Lei (叶磊)
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.
# ----------------------------------------------------------------------

import csv
import os
import random
import re
import traceback

import bs4
import docx
import openpyxl
from PyQt5 import QtCore
from PyQt5 import QtGui
from PyQt5 import QtWidgets

from wordless.wl_checks import (
    wl_checks_misc,
    wl_checks_work_area
)
from wordless.wl_dialogs import (
    wl_dialogs,
    wl_dialogs_misc
)
from wordless.wl_nlp import wl_nlp_utils
from wordless.wl_utils import (
    wl_excs,
    wl_misc,
    wl_paths,
    wl_threading
)
from wordless.wl_widgets import wl_buttons

_tr = QtCore.QCoreApplication.translate

# pylint: disable=unnecessary-lambda

# self.tr() does not work in inherited classes
class Wl_Table(QtWidgets.QTableView):
    def __init__(
        self, parent,
        headers, header_orientation = 'hor',
        editable = False,
        drag_drop = False
    ):
        super().__init__(parent)

        self.main = wl_misc.find_wl_main(parent)

        self.headers = headers
        self.header_orientation = header_orientation

        self.settings = self.main.settings_custom
        self.table_settings = {}

        model = QtGui.QStandardItemModel()
        model.table = self

        self.setModel(model)

        match header_orientation:
            case 'hor':
                self.model().setHorizontalHeaderLabels(self.headers)
            case 'vert':
                self.model().setVerticalHeaderLabels(self.headers)

        self.verticalHeader().setSectionResizeMode(QtWidgets.QHeaderView.ResizeToContents)

        self.horizontalHeader().setHighlightSections(False)
        self.verticalHeader().setHighlightSections(False)

        self.setVerticalScrollMode(QtWidgets.QAbstractItemView.ScrollPerPixel)
        self.setHorizontalScrollMode(QtWidgets.QAbstractItemView.ScrollPerPixel)

        if editable:
            self.setEditTriggers(QtWidgets.QAbstractItemView.DoubleClicked | QtWidgets.QAbstractItemView.SelectedClicked | QtWidgets.QAbstractItemView.EditKeyPressed)
        else:
            self.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)

        if drag_drop:
            self.setDragEnabled(True)
            self.setAcceptDrops(True)
            self.viewport().setAcceptDrops(True)
            self.setDragDropMode(QtWidgets.QAbstractItemView.InternalMove)
            self.setDragDropOverwriteMode(False)

        self.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        # Remove dotted gray border around selected cells
        self.setFocusPolicy(QtCore.Qt.NoFocus)

        self.default_foreground = '#292929'
        self.default_background = '#FFF'

        stylesheet_items = f'''
            QTableView::item:hover {{
                background-color: #E5E5E5;
                color: {self.default_foreground};
            }}
            QTableView::item:selected {{
                background-color: #E5E5E5;
                color: {self.default_foreground};
            }}

            QHeaderView::section {{
                color: {self.default_background};
                font-weight: bold;
            }}
        '''

        match header_orientation:
            case 'hor':
                self.setStyleSheet(stylesheet_items + '''
                    QHeaderView::section:horizontal {
                        background-color: #5C88C5;
                    }
                    QHeaderView::section:horizontal:hover {
                        background-color: #3265B2;
                    }
                    QHeaderView::section:horizontal:pressed {
                        background-color: #264E8C;
                    }

                    QHeaderView::section:vertical {
                        background-color: #888;
                    }
                    QHeaderView::section:vertical:hover {
                        background-color: #777;
                    }
                    QHeaderView::section:vertical:pressed {
                        background-color: #666;
                    }
                ''')
            case 'vert':
                self.setStyleSheet(stylesheet_items + '''
                    QHeaderView::section:horizontal {
                        background-color: #888;
                    }
                    QHeaderView::section:horizontal:hover {
                        background-color: #777;
                    }
                    QHeaderView::section:horizontal:pressed {
                        background-color: #666;
                    }

                    QHeaderView::section:vertical {
                        background-color: #5C88C5;
                    }
                    QHeaderView::section:vertical:hover {
                        background-color: #3265B2;
                    }
                    QHeaderView::section:vertical:pressed {
                        background-color: #264E8C;
                    }
                ''')

        self.model().itemChanged.connect(self.item_changed)
        self.selectionModel().selectionChanged.connect(self.selection_changed)

    # There seems to be a bug with QAbstractItemView.InternalMove
    # See: https://bugreports.qt.io/browse/QTBUG-87057
    def dropEvent(self, event):
        if self.indexAt(event.pos()).row() == -1:
            row_dropped_on = self.model().rowCount()
        else:
            row_dropped_on = self.indexAt(event.pos()).row()

        data = []
        data_selected = []
        rows_selected = self.get_selected_rows()

        for row in range(self.model().rowCount()):
            data.append([])

            for col in range(self.model().columnCount()):
                item = self.model().takeItem(row, col)
                widget = self.indexWidget(self.model().index(row, col))

                if widget:
                    data[-1].append(type(widget)(widget.text(), self))
                else:
                    data[-1].append(item)

        for row in rows_selected:
            data_selected.append(data[row].copy())
            data[row] = []

        for row in reversed(data_selected):
            data.insert(row_dropped_on, row)

        data = [row for row in data if row]

        self.disable_updates()

        self.clr_table()
        self.model().setRowCount(len(data))
        self.model().setColumnCount(len(data[0]))

        for i, row in enumerate(data):
            for j, item in enumerate(row):
                if isinstance(item, QtGui.QStandardItem):
                    self.model().setItem(i, j, item)
                else:
                    self.setIndexWidget(self.model().index(i, j), item)

        self.enable_updates()

        event.accept()

    def item_changed(self):
        if self.is_empty():
            self.setEnabled(False)
        else:
            self.setEnabled(True)

        self.resizeColumnsToContents()
        self.resizeRowsToContents()

        for i in range(self.model().columnCount()):
            self.setColumnWidth(i, self.columnWidth(i) + 22)

        self.selection_changed()

    def selection_changed(self):
        pass

    def disable_updates(self):
        self.num_rows_old = self.model().rowCount()
        self.num_cols_old = self.model().columnCount()

        self.setUpdatesEnabled(False)
        self.blockSignals(True)
        self.horizontalHeader().blockSignals(True)
        self.verticalHeader().blockSignals(True)
        self.model().blockSignals(True)
        self.selectionModel().blockSignals(True)
        self.hide()

    def enable_updates(self, emit_signals = True):
        self.setUpdatesEnabled(True)
        self.blockSignals(False)
        self.horizontalHeader().blockSignals(False)
        self.verticalHeader().blockSignals(False)
        self.model().blockSignals(False)
        self.selectionModel().blockSignals(False)
        self.show()

        if emit_signals:
            self.horizontalHeader().sectionCountChanged.emit(self.num_cols_old, self.model().columnCount())
            self.verticalHeader().sectionCountChanged.emit(self.num_rows_old, self.model().rowCount())

            self.model().itemChanged.emit(QtGui.QStandardItem())
            self.selectionModel().selectionChanged.emit(QtCore.QItemSelection(), QtCore.QItemSelection())

    def is_empty(self):
        match self.header_orientation:
            case 'hor':
                return not any((
                    self.model().item(0, i) or self.indexWidget(self.model().index(0, i))
                    for i in range(self.model().columnCount())
                ))
            case 'vert':
                return not any((
                    self.model().item(i, 0) or self.indexWidget(self.model().index(i, 0))
                    for i in range(self.model().rowCount())
                ))

    def is_visible(self):
        return any((
            not self.isRowHidden(i)
            for i in range(self.model().rowCount())
        ))

    def is_selected(self):
        return bool(self.selectionModel().selectedIndexes())

    def get_header_labels_hor(self):
        return (
            self.model().headerData(row, QtCore.Qt.Horizontal)
            for row in range(self.model().columnCount())
        )

    def get_header_labels_vert(self):
        return (
            self.model().headerData(col, QtCore.Qt.Vertical)
            for col in range(self.model().rowCount())
        )

    def find_header_hor(self, text):
        return list(self.get_header_labels_hor()).index(text)

    def find_header_vert(self, text):
        return list(self.get_header_labels_vert()).index(text)

    def find_headers_hor(self, text):
        return [
            i
            for i, header in enumerate(self.get_header_labels_hor())
            if text in header
        ]

    def find_headers_vert(self, text):
        return [
            i
            for i, header in enumerate(self.get_header_labels_vert())
            if text in header
        ]

    def find_header(self, text):
        match self.header_orientation:
            case 'hor':
                return self.find_header_hor(text = text)
            case 'vert':
                return self.find_header_vert(text = text)

    def find_headers(self, text):
        match self.header_orientation:
            case 'hor':
                return self.find_headers_hor(text = text)
            case 'vert':
                return self.find_headers_vert(text = text)

    def add_header_hor(self, label):
        self.add_headers_hor(labels = [label])

    def add_header_vert(self, label):
        self.add_headers_vert(labels = [label])

    def add_headers_hor(self, labels):
        self.ins_headers_hor(i = self.model().columnCount(), labels = labels)

    def add_headers_vert(self, labels):
        self.ins_headers_vert(i = self.model().rowCount(), labels = labels)

    def ins_header_hor(self, i, label):
        self.ins_headers_hor(i = i, labels = [label])

    def ins_header_vert(self, i, label):
        self.ins_headers_vert(i = i, labels = [label])

    def ins_headers_hor(self, i, labels):
        headers = list(self.get_header_labels_hor())
        headers[i:i] = labels

        self.model().setHorizontalHeaderLabels(headers)

    def ins_headers_vert(self, i, labels):
        headers = list(self.get_header_labels_vert())
        headers[i:i] = labels

        self.model().setVerticalHeaderLabels(headers)

    def get_selected_rows(self, visible_only = False):
        selected_rows = sorted({index.row() for index in self.selectionModel().selectedIndexes()})

        if visible_only:
            return [row for row in selected_rows if not self.isRowHidden(row)]
        else:
            return selected_rows

    def get_selected_cols(self, visible_only = False):
        selected_col = sorted({index.column() for index in self.selectionModel().selectedIndexes()})

        if visible_only:
            return [col for col in selected_col if not self.isColumnHidden(col)]
        else:
            return selected_col

    def _add_row(self, row = None, texts = None):
        if texts is None:
            texts = self.defaults_row

        if self.is_empty():
            self.clr_table(0)

        if row is None:
            self.model().appendRow([QtGui.QStandardItem(text) for text in texts])
        else:
            self.model().insertRow(row, [QtGui.QStandardItem(text) for text in texts])

        self.model().itemChanged.emit(QtGui.QStandardItem())

    def add_row(self, texts = None):
        self._add_row(texts = texts)
        self.setCurrentIndex(self.model().index(self.model().rowCount() - 1, 0))

    def ins_row(self, texts = None):
        row = self.get_selected_rows()[0]

        self._add_row(row = row, texts = texts)
        self.setCurrentIndex(self.model().index(row, 0))

    def del_row(self):
        for i in reversed(self.get_selected_rows()):
            self.model().removeRow(i)

        self.model().itemChanged.emit(QtGui.QStandardItem())

    def clr_table(self, num_headers = 1):
        self.model().clear()

        match self.header_orientation:
            case 'hor':
                self.model().setHorizontalHeaderLabels(self.headers)
                self.model().setRowCount(num_headers)
            case 'vert':
                self.model().setVerticalHeaderLabels(self.headers)
                self.model().setColumnCount(num_headers)

        self.model().itemChanged.emit(QtGui.QStandardItem())

    # Export visible rows only
    @wl_misc.log_time
    def exp_selected_cells(self):
        return self.exp_all_cells(rows_to_exp = self.get_selected_rows(visible_only = True))

    @wl_misc.log_time
    def exp_all_cells(self, rows_to_exp = None):
        caption = _tr('wl_tables', 'Export Table')
        default_dir = self.main.settings_custom['general']['exp']['tables']['default_path']
        default_type = self.main.settings_custom['general']['exp']['tables']['default_type']
        default_ext = re.search(r'(?<=\(\*\.)[a-zA-Z0-9]+(?=[;\)])', default_type).group()

        # Errors (Search terms, stop word lists, file checking, etc.)
        match self.tab:
            case 'err':
                file_path, file_type = QtWidgets.QFileDialog.getSaveFileName(
                    parent = self,
                    caption = caption,
                    directory = os.path.join(wl_checks_misc.check_dir(default_dir), f'wordless_error.{default_ext}'),
                    filter = ';;'.join(self.main.settings_global['file_types']['exp_tables']),
                    initialFilter = default_type
                )
            case 'concordancer' | 'concordancer_parallel':
                # Concordancer (with zapping)
                if self.tab == 'concordancer' and self.main.settings_custom['concordancer']['zapping_settings']['zapping']:
                    file_path, file_type = QtWidgets.QFileDialog.getSaveFileName(
                        parent = self,
                        caption = caption,
                        directory = os.path.join(wl_checks_misc.check_dir(default_dir), f'wordless_results_{self.tab}.docx'),
                        filter = ';;'.join(self.main.settings_global['file_types']['exp_tables_concordancer_zapping']),
                    )
                # Concordancer (without zapping) & Parallel Concordancer
                else:
                    file_path, file_type = QtWidgets.QFileDialog.getSaveFileName(
                        parent = self,
                        caption = caption,
                        directory = os.path.join(wl_checks_misc.check_dir(default_dir), f'wordless_results_{self.tab}.{default_ext}'),
                        filter = ';;'.join(self.main.settings_global['file_types']['exp_tables_concordancer']),
                        initialFilter = default_type
                    )
            # Other modules
            case _:
                file_path, file_type = QtWidgets.QFileDialog.getSaveFileName(
                    parent = self,
                    caption = caption,
                    directory = os.path.join(wl_checks_misc.check_dir(default_dir), f'wordless_results_{self.tab}.{default_ext}'),
                    filter = ';;'.join(self.main.settings_global['file_types']['exp_tables']),
                    initialFilter = default_type
                )

        if file_path:
            self.worker_exp_table = Wl_Worker_Exp_Table(
                self.main,
                dialog_progress = wl_dialogs_misc.Wl_Dialog_Progress(
                    self.main,
                    text = _tr('wl_tables', 'Exporting table...')
                ),
                table = self,
                file_path = file_path,
                file_type = file_type,
                rows_to_exp = (
                    rows_to_exp or
                    # Export visible rows only
                    [row for row in range(self.model().rowCount()) if not self.isRowHidden(row)]
                )
            )

            self.thread_exp_table = QtCore.QThread()
            wl_threading.start_worker_in_thread(
                self.worker_exp_table,
                self.thread_exp_table,
                self.update_gui_exp
            )

            return ''
        # Do not log time if the export dialog is closed
        else:
            return 'skip_logging_time'

    def update_gui_exp(self, err_msg, file_path):
        if not err_msg:
            self.results_saved = True

        wl_checks_work_area.check_err_exp_table(self.parent(), err_msg, file_path)

RE_REDUNDANT_SPACES = re.compile(r'\s+')
RE_INVALID_XML_CHARS = re.compile(r'[^\u0009\u000A\u000D\u0020-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]+')
RE_COLOR = re.compile(r'(?<=color: #)([0-9a-fA-F]{3}|[0-9a-fA-F]{6})(?=;)')

class Wl_Worker_Exp_Table(wl_threading.Wl_Worker):
    finished = QtCore.pyqtSignal(str, str)

    def run(self):
        err_msg = ''

        try:
            if 'headers_int' not in self.table.__dict__:
                self.table.headers_int = set()
            if 'headers_float' not in self.table.__dict__:
                self.table.headers_float = set()
            if 'headers_pct' not in self.table.__dict__:
                self.table.headers_pct = set()

            settings_concordancer = self.main.settings_custom['concordancer']['zapping_settings']

            num_rows = len(self.rows_to_exp)
            # Export visible columns only
            cols = [col for col in range(self.table.model().columnCount()) if not self.table.isColumnHidden(col)]

            # CSV files
            if '*.csv' in self.file_type:
                encoding = self.main.settings_custom['general']['exp']['tables']['default_encoding']

                with open(self.file_path, 'w', encoding = encoding, newline = '') as f:
                    csv_writer = csv.writer(f)

                    if self.table.header_orientation == 'hor':
                        # Horizontal headers
                        headers_hor = [
                            self.table.model().horizontalHeaderItem(col).text()
                            for col in cols
                        ]
                        csv_writer.writerow(self.clean_text_csv(headers_hor))

                        # Cells
                        for i, row in enumerate(self.rows_to_exp):
                            self.progress_updated.emit(self.tr('Exporting table... ({} / {})').format(i + 1, num_rows))

                            row_to_exp = []

                            for col in cols:
                                if not self._running:
                                    raise wl_excs.Wl_Exc_Aborted(self.main)

                                if self.table.model().item(row, col):
                                    cell_text = self.table.model().item(row, col).text()
                                else:
                                    cell_text = self.table.indexWidget(self.table.model().index(row, col)).text()
                                    cell_text = wl_nlp_utils.html_to_text(cell_text)

                                row_to_exp.append(cell_text)

                            csv_writer.writerow(self.clean_text_csv(row_to_exp))
                    # Profiler
                    else:
                        # Horizontal headers
                        headers_hor = [
                            self.table.model().horizontalHeaderItem(col).text()
                            for col in cols
                        ]
                        csv_writer.writerow([''] + self.clean_text_csv(headers_hor))

                        # Vertical headers and cells
                        for i, row in enumerate(self.rows_to_exp):
                            self.progress_updated.emit(self.tr('Exporting table... ({} / {})').format(i + 1, num_rows))

                            row_to_exp = [self.table.model().verticalHeaderItem(row).text()]

                            for col in cols:
                                if not self._running:
                                    raise wl_excs.Wl_Exc_Aborted(self.main)

                                row_to_exp.append(self.table.model().item(row, col).text())

                            csv_writer.writerow(self.clean_text_csv(row_to_exp))
            # Excel workbooks
            elif '*.xlsx' in self.file_type:
                workbook = openpyxl.Workbook()
                worksheet = workbook.active

                dpi_horizontal = QtWidgets.QApplication.primaryScreen().logicalDotsPerInchX()
                dpi_vertical = QtWidgets.QApplication.primaryScreen().logicalDotsPerInchY()

                match self.table.tab:
                    case 'concordancer':
                        freeze_panes = 'A2'

                        # Left, Node, Right
                        cols_labels = [0, 1, 2]
                        cols_table_items = []
                    case 'concordancer_parallel':
                        freeze_panes = 'A2'

                        cols_labels = []
                        # Parallel Unit No. (%)
                        cols_table_items = [0, 1]
                    case 'dependency_parser':
                        freeze_panes = 'A2'

                        # Sentence
                        cols_labels = [5]
                        cols_table_items = []
                    case _:
                        freeze_panes = 'B2'

                        cols_labels = []
                        cols_table_items = []

                worksheet.freeze_panes = freeze_panes

                if self.table.header_orientation == 'hor':
                    # Horizontal headers
                    for col_cell, col_item in enumerate(cols):
                        if not self._running:
                            raise wl_excs.Wl_Exc_Aborted(self.main)

                        cell = worksheet.cell(1, 1 + col_cell)
                        cell.value = self.table.model().horizontalHeaderItem(col_item).text()

                        self.style_header_hor(cell)

                        worksheet.column_dimensions[openpyxl.utils.get_column_letter(1 + col_cell)].width = self.table.horizontalHeader().sectionSize(col_item) / dpi_horizontal * 13 + 3

                    # Cells
                    for row_cell, row_item in enumerate(self.rows_to_exp):
                        self.progress_updated.emit(self.tr('Exporting table... ({} / {})').format(row_cell + 1, num_rows))

                        for col_cell, col_item in enumerate(cols):
                            if not self._running:
                                raise wl_excs.Wl_Exc_Aborted(self.main)

                            cell = worksheet.cell(2 + row_cell, 1 + col_cell)

                            if (
                                (
                                    cols_labels
                                    and not cols_table_items
                                    and col_item in cols_labels
                                ) or (
                                    not cols_labels
                                    and cols_table_items
                                    and col_item not in cols_table_items
                                )
                            ):
                                cell_val = self.table.indexWidget(self.table.model().index(row_item, col_item)).text()
                                cell_val = self.remove_invalid_xml_chars(cell_val)
                                cell.value = cell_val

                                self.style_cell_rich_text(cell, self.table.indexWidget(self.table.model().index(row_item, col_item)))
                            else:
                                cell_val = self.table.model().item(row_item, col_item).text()
                                cell_val = self.remove_invalid_xml_chars(cell_val)
                                cell.value = cell_val

                                self.style_cell(cell, self.table.model().item(row_item, col_item))
                # Profiler
                else:
                    # Horizontal headers
                    for col_cell, col_item in enumerate(cols):
                        if not self._running:
                            raise wl_excs.Wl_Exc_Aborted(self.main)

                        cell = worksheet.cell(1, 2 + col_cell)
                        cell.value = self.table.model().horizontalHeaderItem(col_item).text()

                        self.style_header_hor(cell)

                        worksheet.column_dimensions[openpyxl.utils.get_column_letter(2 + col_cell)].width = self.table.horizontalHeader().sectionSize(col_item) / dpi_horizontal * 13 + 3

                    worksheet.column_dimensions[openpyxl.utils.get_column_letter(1)].width = self.table.verticalHeader().width() / dpi_horizontal * 13 + 3

                    # Vertical headers
                    for row_cell, row_item in enumerate(self.rows_to_exp):
                        if not self._running:
                            raise wl_excs.Wl_Exc_Aborted(self.main)

                        cell = worksheet.cell(2 + row_cell, 1)
                        cell.value = self.table.model().verticalHeaderItem(row_item).text()

                        self.style_header_vert(cell)

                    # Cells
                    for row_cell, row_item in enumerate(self.rows_to_exp):
                        self.progress_updated.emit(self.tr('Exporting table... ({} / {})').format(row_cell + 1, num_rows))

                        for col_cell, col_item in enumerate(cols):
                            if not self._running:
                                raise wl_excs.Wl_Exc_Aborted(self.main)

                            cell = worksheet.cell(2 + row_cell, 2 + col_cell)
                            cell_val = self.table.model().item(row_item, col_item).text()
                            cell_val = self.remove_invalid_xml_chars(cell_val)
                            cell.value = cell_val

                            self.style_cell(cell, self.table.model().item(row_item, col_item))

                # Row height
                worksheet.row_dimensions[1].height = self.table.horizontalHeader().height() / dpi_vertical * 72

                for i, _ in enumerate(worksheet.rows):
                    worksheet.row_dimensions[2 + i].height = self.table.verticalHeader().sectionSize(0) / dpi_vertical * 72

                self.progress_updated.emit(self.tr('Saving file...'))

                workbook.save(self.file_path)
            elif '*.docx' in self.file_type:
                doc = docx.Document()

                # Concordancer
                if self.table.tab == 'concordancer':
                    outputs = []

                    for i, row in enumerate(self.rows_to_exp):
                        if not self._running:
                            raise wl_excs.Wl_Exc_Aborted(self.main)

                        self.progress_updated.emit(self.tr('Processing data... ({} / {})').format(i + 1, num_rows))

                        para_text = []

                        for col in range(3):
                            para_text.append(self.table.indexWidget(self.table.model().index(row, col)).text().strip())

                        # Zapping
                        if settings_concordancer['zapping']:
                            # Node
                            para_text[1] = settings_concordancer['placeholder'] * settings_concordancer['replace_keywords_with']

                        outputs.append([' '.join(para_text), self.table.indexWidget(self.table.model().index(row, col))])

                    if settings_concordancer['zapping']:
                        # Randomize outputs
                        if settings_concordancer['randomize_outputs']:
                            random.shuffle(outputs)

                        # Assign line numbers
                        if settings_concordancer['add_line_nums']:
                            for i, _ in enumerate(outputs):
                                if not self._running:
                                    raise wl_excs.Wl_Exc_Aborted(self.main)

                                outputs[i][0] = f'{i + 1}. ' + outputs[i][0]

                    for i, (para_text, item) in enumerate(outputs):
                        if not self._running:
                            raise wl_excs.Wl_Exc_Aborted(self.main)

                        self.progress_updated.emit(self.tr('Exporting table... ({} / {})').format(i + 1, num_rows))

                        para = self.add_para(doc)
                        self.style_para_rich_text(para, para_text, item)
                # Parallel Concordancer
                elif self.table.tab == 'concordancer_parallel':
                    for i, row in enumerate(self.rows_to_exp):
                        self.progress_updated.emit(self.tr('Exporting table... ({} / {})').format(i + 1, num_rows))

                        if i > 0:
                            self.add_para(doc)

                        for col in range(2, self.table.model().columnCount()):
                            if not self._running:
                                raise wl_excs.Wl_Exc_Aborted(self.main)

                            para_text = self.table.indexWidget(self.table.model().index(row, col)).text().strip()

                            para = self.add_para(doc)
                            self.style_para_rich_text(para, para_text, self.table.indexWidget(self.table.model().index(row, col)))

                # Add the last empty paragraph
                self.add_para(doc)

                self.progress_updated.emit(self.tr('Saving file...'))

                doc.save(self.file_path)
        except wl_excs.Wl_Exc_Aborted:
            if '*.csv' in self.file_type:
                if os.path.exists(self.file_path):
                    os.remove(self.file_path)

            err_msg = 'aborted'
        except PermissionError:
            err_msg = 'permission_err'
        except Exception: # pylint: disable=broad-exception-caught
            err_msg = traceback.format_exc()

        self.main.settings_custom['general']['exp']['tables']['default_path'] = wl_paths.get_normalized_dir(self.file_path)
        self.main.settings_custom['general']['exp']['tables']['default_type'] = self.file_type

        self.finished.emit(err_msg, self.file_path)

    # Clean text before writing to CSV files
    def clean_text_csv(self, items):
        for i, item in enumerate(items):
            items[i] = item.replace('\n', ' ')
            items[i] = RE_REDUNDANT_SPACES.sub(' ', items[i])
            items[i] = items[i].strip()

        return items

    # Remove invalid XML characters
    def remove_invalid_xml_chars(self, text):
        # openpyxl.cell.cell.ILLEGAL_CHARACTERS_RE is not complete
        # Reference: https://www.w3.org/TR/xml/#charsets
        return RE_INVALID_XML_CHARS.sub('', text)

    def style_header(self, cell):
        cell.font = openpyxl.styles.Font(
            name = self.main.settings_custom['general']['ui_settings']['font_family'],
            size = self.main.settings_custom['general']['ui_settings']['font_size'],
            bold = True,
            color = 'FFFFFF'
        )

    def style_header_hor(self, cell):
        self.style_header(cell)

        match self.table.header_orientation:
            # Headers
            case 'hor':
                cell.fill = openpyxl.styles.PatternFill(
                    fill_type = 'solid',
                    fgColor = '5C88C5'
                )
            # File names
            case 'vert':
                cell.fill = openpyxl.styles.PatternFill(
                    fill_type = 'solid',
                    fgColor = '888888'
                )

        cell.alignment = openpyxl.styles.Alignment(
            horizontal = 'center',
            vertical = 'center',
            wrap_text = True
        )

    def style_header_vert(self, cell):
        self.style_header(cell)

        match self.table.header_orientation:
            # Line numbers
            case 'hor':
                cell.fill = openpyxl.styles.PatternFill(
                    fill_type = 'solid',
                    fgColor = '888888'
                )
                cell.alignment = openpyxl.styles.Alignment(
                    horizontal = 'right',
                    vertical = 'center',
                    wrap_text = True
                )
            # Headers
            case 'vert':
                cell.fill = openpyxl.styles.PatternFill(
                    fill_type = 'solid',
                    fgColor = '5C88C5'
                )
                cell.alignment = openpyxl.styles.Alignment(
                    horizontal = 'left',
                    vertical = 'center',
                    wrap_text = True
                )

    def style_cell_alignment(self, cell, item):
        if isinstance(item, QtGui.QStandardItem):
            alignment = item.textAlignment()
        else:
            alignment = item.alignment()

        if alignment & QtCore.Qt.AlignLeft == QtCore.Qt.AlignLeft:
            alignment_hor = 'left'
        elif alignment & QtCore.Qt.AlignRight == QtCore.Qt.AlignRight:
            alignment_hor = 'right'
        elif alignment & QtCore.Qt.AlignHCenter == QtCore.Qt.AlignHCenter:
            alignment_hor = 'center'
        elif alignment & QtCore.Qt.AlignJustify == QtCore.Qt.AlignJustify:
            alignment_hor = 'justify'
        # Default
        else:
            alignment_hor = 'left'

        if alignment & QtCore.Qt.AlignTop == QtCore.Qt.AlignTop:
            alignment_vert = 'top'
        elif alignment & QtCore.Qt.AlignBottom == QtCore.Qt.AlignBottom:
            alignment_vert = 'bottom'
        elif alignment & QtCore.Qt.AlignVCenter == QtCore.Qt.AlignVCenter:
            alignment_vert = 'center'
        # Not sure
        elif alignment & QtCore.Qt.AlignBaseline == QtCore.Qt.AlignBaseline:
            alignment_vert = 'justify'
        # Default
        else:
            alignment_vert = 'center'

        cell.alignment = openpyxl.styles.Alignment(
            horizontal = alignment_hor,
            vertical = alignment_vert,
            wrap_text = True
        )

    def style_cell(self, cell, item):
        # Modify number format
        val = cell.value

        try:
            if val[-1] == '%':
                cell.value = float(val[:-1]) / 100
            else:
                cell.value = float(val)

            if val[-1] == '%':
                precision_pcts = self.main.settings_custom['tables']['precision_settings']['precision_pcts']

                if precision_pcts:
                    cell.number_format = '0.' + '0' * precision_pcts + '%'
                else:
                    cell.number_format = '0%'
            else:
                i_decimal_point = val.find('.')

                if i_decimal_point > -1:
                    cell.number_format = '0.' + '0' * (len(val) - i_decimal_point - 1)
                else:
                    cell.number_format = '0'
        # Skip text
        except ValueError:
            pass

        font_family = item.font().family()

        if font_family != 'Consolas':
            font_family = self.main.settings_custom['general']['ui_settings']['font_family']

        cell.font = openpyxl.styles.Font(
            name = font_family,
            size = self.main.settings_custom['general']['ui_settings']['font_size'],
            bold = item.font().bold(),
            italic = item.font().italic()
        )

        self.style_cell_alignment(cell, item)

    def style_cell_rich_text(self, cell, item):
        rich_texts = []
        font_family = item.font().family()

        if font_family != 'Consolas':
            font_family = self.main.settings_custom['general']['ui_settings']['font_family']

        # Wrap HTML with <html><body><p></p></body></html>
        soup = bs4.BeautifulSoup(f'<html><body><p>{cell.value}</p></body></html>', features = 'lxml')

        for html in soup.body.p.contents:
            if isinstance(html, bs4.element.Tag):
                text = html.text.strip()
            else:
                text = html.strip()

            if text:
                if isinstance(html, bs4.element.Tag) and html.has_attr('style'):
                    style = html['style']

                    if (re_color := RE_COLOR.search(style)):
                        color = re_color.group()

                        # 3-digit color shorthand
                        if len(color) == 3:
                            color = color[0] * 2 + color[1] * 2 + color[2] * 2
                    else:
                        color = '000000'

                    bold = 'font-weight: bold;' in style
                    italic = 'font-style: italic;' in style

                    rich_texts.append(openpyxl.cell.rich_text.TextBlock(
                        openpyxl.cell.text.InlineFont(
                            rFont = font_family,
                            sz = self.main.settings_custom['general']['ui_settings']['font_size'],
                            b = bold,
                            i = italic,
                            # 6/8-digit aRGB hex values without "#"
                            color = color
                        ),
                        text + ' '
                    ))
                else:
                    rich_texts.append(openpyxl.cell.rich_text.TextBlock(
                        openpyxl.cell.text.InlineFont(
                            rFont = font_family,
                            sz = self.main.settings_custom['general']['ui_settings']['font_size'],
                            color = '000000'
                        ),
                        text + ' '
                    ))

        if rich_texts:
            # Remove trailing space after the last part of the text
            rich_texts[-1].text = rich_texts[-1].text.strip()

        cell.value = openpyxl.cell.rich_text.CellRichText(rich_texts)

        self.style_cell_alignment(cell, item)

    def style_cell_concordancer_node(self, cell, item):
        cell.font = openpyxl.styles.Font(
            name = item.font().family(),
            size = self.main.settings_custom['general']['ui_settings']['font_size'],
            bold = True,
            color = 'FF0000'
        )

        self.style_cell_alignment(cell, item)

    def add_para(self, doc):
        para = doc.add_paragraph()
        para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

        self.style_para_spacing(para)

        return para

    def style_para_rich_text(self, para, para_text, item):
        font_family = item.font().family()

        if font_family != 'Consolas':
            font_family = self.main.settings_custom['general']['ui_settings']['font_family']

        # Wrap HTML with <html><body><p></p></body></html>
        soup = bs4.BeautifulSoup(f'<html><body><p>{para_text}</p></body></html>', features = 'lxml')

        for html in soup.body.p.contents:
            if isinstance(html, bs4.element.Tag):
                text = html.text.strip()
            else:
                text = html.strip()

            if text:
                if para.text:
                    para.add_run(' ')

                run = para.add_run(text)

                run.font.name = self.main.settings_custom['general']['ui_settings']['font_family']
                run.font.size = docx.shared.Pt(self.main.settings_custom['general']['ui_settings']['font_size'])

                if isinstance(html, bs4.element.Tag) and html.has_attr('style'):
                    style = html['style']

                    if (re_color := RE_COLOR.search(style)):
                        color = re_color.group()

                        # 3-digit color shorthand
                        if len(color) == 3:
                            color = color[0] * 2 + color[1] * 2 + color[2] * 2
                    else:
                        color = '000000'

                    bold = 'font-weight: bold;' in style
                    italic = 'font-style: italic;' in style

                    run.bold = bold
                    run.italic = italic
                    run.font.color.rgb = docx.shared.RGBColor.from_string(color)
                else:
                    run.font.color.rgb = docx.shared.RGBColor.from_string('000000')

        self.style_para_spacing(para)

    def style_para_spacing(self, para):
        para.paragraph_format.space_before = docx.shared.Pt(0)
        para.paragraph_format.space_after = docx.shared.Pt(0)
        para.paragraph_format.line_spacing = 1.5

class Wl_Table_Add_Ins_Del_Clr(Wl_Table):
    def __init__(self, parent, headers, col_edit = None):
        super().__init__(
            parent = parent,
            headers = headers,
            editable = True,
            drag_drop = True
        )

        self.col_edit = col_edit

        self.button_add = QtWidgets.QPushButton(_tr('wl_tables', 'Add'), self)
        self.button_ins = QtWidgets.QPushButton(_tr('wl_tables', 'Insert'), self)
        self.button_del = QtWidgets.QPushButton(_tr('wl_tables', 'Remove'), self)
        self.button_clr = QtWidgets.QPushButton(_tr('wl_tables', 'Clear'), self)

        self.button_add.clicked.connect(lambda: self.add_row())
        self.button_ins.clicked.connect(lambda: self.ins_row())
        self.button_del.clicked.connect(lambda: self.del_row())
        self.button_clr.clicked.connect(lambda: self.clr_table(0))

    def item_changed(self):
        if not self.is_empty():
            self.button_clr.setEnabled(True)
        else:
            self.button_clr.setEnabled(False)

        super().item_changed()

    def selection_changed(self):
        if self.selectionModel().selectedIndexes():
            self.button_ins.setEnabled(True)
            self.button_del.setEnabled(True)
        else:
            self.button_ins.setEnabled(False)
            self.button_del.setEnabled(False)

    def add_row(self, texts = None):
        super().add_row(texts = texts)

        if self.col_edit is not None:
            self.edit(self.model().index(self.model().rowCount() - 1, self.col_edit))

    def ins_row(self, texts = None):
        super().ins_row(texts = texts)

        if self.col_edit is not None:
            self.edit(self.model().index(self.get_selected_rows()[0], self.col_edit))

class Wl_Table_Item(QtGui.QStandardItem):
    def read_data(self):
        if (
            self.column() in self.model().table.headers_int
            or self.column() in self.model().table.headers_float
            or self.column() in self.model().table.headers_pct
        ):
            return self.val
        else:
            return self.text()

    def __lt__(self, other):
        return self.read_data() < other.read_data()

class Wl_Table_Item_Err(QtGui.QStandardItem):
    def read_data(self):
        return self.text()

    def __lt__(self, other):
        return self.read_data() < other.read_data()

class Wl_Table_Data(Wl_Table):
    def __init__(
        self, main, tab,
        headers, header_orientation = 'hor',
        headers_int = None, headers_float = None,
        headers_pct = None, headers_cum = None,
        cols_breakdown_file = None, cols_breakdown_span_position = None,
        enable_sorting = False, generate_fig = True
    ):
        super().__init__(
            main, headers, header_orientation,
            editable = False,
            drag_drop = False
        )

        self.tab = tab

        self.headers_int_old = headers_int or set()
        self.headers_float_old = headers_float or set()
        self.headers_pct_old = headers_pct or set()
        self.headers_cum_old = headers_cum or set()
        self.cols_breakdown_file_old = cols_breakdown_file or set()
        self.cols_breakdown_span_position_old = cols_breakdown_span_position or set()

        self.enable_sorting = enable_sorting

        if enable_sorting:
            self.setSortingEnabled(True)

            match header_orientation:
                case 'hor':
                    self.horizontalHeader().sortIndicatorChanged.connect(self.sorting_changed)
                case 'vert':
                    self.verticalHeader().sortIndicatorChanged.connect(self.sorting_changed)

        self.model().itemChanged.connect(self.item_changed)
        self.selectionModel().selectionChanged.connect(self.selection_changed)

        self.button_generate_table = QtWidgets.QPushButton(_tr('wl_tables', 'Generate table'), self)
        self.button_generate_fig = QtWidgets.QPushButton(_tr('wl_tables', 'Generate figure'), self)
        self.button_exp_selected_cells = QtWidgets.QPushButton(_tr('wl_tables', 'Export selected cells...'), self)
        self.button_exp_all_cells = QtWidgets.QPushButton(_tr('wl_tables', 'Export all cells...'), self)
        self.button_clr_table = QtWidgets.QPushButton(_tr('wl_tables', 'Clear table'), self)

        if not generate_fig:
            self.button_generate_fig.hide()

        self.button_generate_table.clicked.connect(lambda: self.generate_table())
        self.button_generate_fig.clicked.connect(lambda: self.generate_fig())
        self.button_exp_selected_cells.clicked.connect(lambda: self.exp_selected_cells())
        self.button_exp_all_cells.clicked.connect(lambda: self.exp_all_cells())
        self.button_clr_table.clicked.connect(lambda: self.clr_table(confirm = True))

        self.main.wl_file_area.table_files.model().itemChanged.connect(self.file_changed)

        self.clr_table()
        self.file_changed()

    def item_changed(self):
        if not self.is_empty() and self.is_visible():
            self.button_exp_all_cells.setEnabled(True)
        else:
            self.button_exp_all_cells.setEnabled(False)

        if not self.is_empty():
            self.button_clr_table.setEnabled(True)
        else:
            self.button_clr_table.setEnabled(False)

        super().item_changed()

        self.selectionModel().selectionChanged.emit(QtCore.QItemSelection(), QtCore.QItemSelection())

    def selection_changed(self):
        # Enable "Export selected cells" only if any visible rows are selected
        if not self.is_empty() and self.get_selected_rows(visible_only = True):
            self.button_exp_selected_cells.setEnabled(True)
        else:
            self.button_exp_selected_cells.setEnabled(False)

    def sorting_changed(self):
        if not self.is_empty():
            if _tr('wl_tables', 'Rank') in self.get_header_labels_hor():
                self.update_ranks()

            if self.table_settings['show_cum_data']:
                self.toggle_cum_data()

    def file_changed(self):
        if list(self.main.wl_file_area.get_selected_files()):
            self.button_generate_table.setEnabled(True)
            self.button_generate_fig.setEnabled(True)
        else:
            self.button_generate_table.setEnabled(False)
            self.button_generate_fig.setEnabled(False)

    def add_header_hor(
        self, label,
        is_int = False, is_float = False,
        is_pct = False, is_cum = False,
        is_breakdown_file = False, is_breakdown_span_position = False
    ):
        self.add_headers_hor(
            labels = [label],
            is_int = is_int, is_float = is_float,
            is_pct = is_pct, is_cum = is_cum,
            is_breakdown_file = is_breakdown_file, is_breakdown_span_position = is_breakdown_span_position
        )

    def add_header_vert(
        self, label,
        is_int = False, is_float = False,
        is_pct = False, is_cum = False
    ):
        self.add_headers_vert(
            labels = [label],
            is_int = is_int, is_float = is_float,
            is_pct = is_pct, is_cum = is_cum
        )

    def add_headers_hor(
        self, labels,
        is_int = False, is_float = False,
        is_pct = False, is_cum = False,
        is_breakdown_file = False, is_breakdown_span_position = False
    ):
        self.ins_headers_hor(
            i = self.model().columnCount(), labels = labels,
            is_int = is_int, is_float = is_float,
            is_pct = is_pct, is_cum = is_cum,
            is_breakdown_file = is_breakdown_file, is_breakdown_span_position = is_breakdown_span_position
        )

    def add_headers_vert(
        self, labels,
        is_int = False, is_float = False,
        is_pct = False, is_cum = False
    ):
        self.ins_headers_vert(
            i = self.model().rowCount(), labels = labels,
            is_int = is_int, is_float = is_float,
            is_pct = is_pct, is_cum = is_cum,
        )

    def ins_header_hor(
        self, i, label,
        is_int = False, is_float = False,
        is_pct = False, is_cum = False,
        is_breakdown_file = False, is_breakdown_span_position = False
    ):
        self.ins_headers_hor(
            i = i, labels = [label],
            is_int = is_int, is_float = is_float,
            is_pct = is_pct, is_cum = is_cum,
            is_breakdown_file = is_breakdown_file, is_breakdown_span_position = is_breakdown_span_position
        )

    def ins_header_vert(
        self, i, label,
        is_int = False, is_float = False,
        is_pct = False, is_cum = False
    ):
        self.ins_headers_vert(
            i = i, labels = [label],
            is_int = is_int, is_float = is_float,
            is_pct = is_pct, is_cum = is_cum
        )

    def ins_headers_hor(
        self, i, labels,
        is_int = False, is_float = False,
        is_pct = False, is_cum = False,
        is_breakdown_file = False, is_breakdown_span_position = False
    ):
        # Re-calculate column indexes
        if self.header_orientation == 'hor':
            headers_int = [
                self.model().horizontalHeaderItem(col).text()
                for col in self.headers_int
            ]
            headers_float = [
                self.model().horizontalHeaderItem(col).text()
                for col in self.headers_float
            ]
            headers_pct = [
                self.model().horizontalHeaderItem(col).text()
                for col in self.headers_pct
            ]
            headers_cum = [
                self.model().horizontalHeaderItem(col).text()
                for col in self.headers_cum
            ]

        cols_breakdown_file = [
            self.model().horizontalHeaderItem(col).text()
            for col in self.cols_breakdown_file
        ]
        cols_breakdown_span_position = [
            self.model().horizontalHeaderItem(col).text()
            for col in self.cols_breakdown_span_position
        ]

        super().ins_headers_hor(i, labels)

        if self.header_orientation == 'hor':
            if is_int:
                headers_int.extend(labels)
            if is_float:
                headers_float.extend(labels)
            if is_pct:
                headers_pct.extend(labels)
            if is_cum:
                headers_cum.extend(labels)

            self.headers_int = {self.find_header_hor(header) for header in headers_int}
            self.headers_float = {self.find_header_hor(header) for header in headers_float}
            self.headers_pct = {self.find_header_hor(header) for header in headers_pct}
            self.headers_cum = {self.find_header_hor(header) for header in headers_cum}

        if is_breakdown_file:
            cols_breakdown_file.extend(labels)
        if is_breakdown_span_position:
            cols_breakdown_span_position.extend(labels)

        self.cols_breakdown_file = {
            self.find_header_hor(header)
            for header in cols_breakdown_file
        }
        self.cols_breakdown_span_position = {
            self.find_header_hor(header)
            for header in cols_breakdown_span_position
        }

    def ins_headers_vert(
        self, i, labels,
        is_int = False, is_float = False,
        is_pct = False, is_cum = False
    ):
        # Re-calculate row indexes
        headers_int = [self.model().verticalHeaderItem(row).text() for row in self.headers_int]
        headers_float = [self.model().verticalHeaderItem(row).text() for row in self.headers_float]
        headers_pct = [self.model().verticalHeaderItem(row).text() for row in self.headers_pct]
        headers_cum = [self.model().verticalHeaderItem(row).text() for row in self.headers_cum]

        super().ins_headers_vert(i, labels)

        if is_int:
            headers_int.extend(labels)
        if is_float:
            headers_float.extend(labels)
        if is_pct:
            headers_pct.extend(labels)
        if is_cum:
            headers_cum.extend(labels)

        self.headers_int = {self.find_header_vert(header) for header in headers_int}
        self.headers_float = {self.find_header_vert(header) for header in headers_float}
        self.headers_pct = {self.find_header_vert(header) for header in headers_pct}
        self.headers_cum = {self.find_header_vert(header) for header in headers_cum}

    def set_item_num(self, row, col, val, total = -1):
        match self.header_orientation:
            case 'hor':
                header = col
            case 'vert':
                header = row

        # Integers
        if header in self.headers_int:
            val = int(val)

            item = Wl_Table_Item(str(val))
        # Floats
        elif header in self.headers_float:
            val = float(val)
            precision = self.main.settings_custom['tables']['precision_settings']['precision_decimals']

            item = Wl_Table_Item(f'{val:.{precision}f}')
        # Percentages
        elif header in self.headers_pct:
            if total > 0:
                val = val / total
            # Handle zero division error
            elif total == 0:
                val = 0
            # Set values directly
            elif total == -1:
                val = float(val)

            precision = self.main.settings_custom['tables']['precision_settings']['precision_pcts']

            item = Wl_Table_Item(f'{val:.{precision}%}')

        item.val = val

        item.setFont(QtGui.QFont('Consolas'))
        item.setTextAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignVCenter)

        self.model().setItem(row, col, item)

    def set_item_num_val(self, row, col, val):
        match self.header_orientation:
            case 'hor':
                header = col
            case 'vert':
                header = row

        item = self.model().item(row, col)

        # Integers
        if header in self.headers_int:
            item.setText(str(val))
        # Floats
        elif header in self.headers_float:
            val = float(val)
            precision = self.main.settings_custom['tables']['precision_settings']['precision_decimals']

            item.setText(f'{val:.{precision}f}')
        # Percentages
        elif header in self.headers_pct:
            precision = self.main.settings_custom['tables']['precision_settings']['precision_pcts']

            item.setText(f'{val:.{precision}%}')

        item.val = val

    def set_item_p_val(self, row, col, val):
        precision = self.main.settings_custom['tables']['precision_settings']['precision_p_vals']
        item = Wl_Table_Item(f'{val:.{precision}f}')

        item.val = val

        item.setFont(QtGui.QFont('Consolas'))
        item.setTextAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignVCenter)

        self.model().setItem(row, col, item)

    def set_item_err(self, row, col, text, alignment_hor = 'center'):
        item = Wl_Table_Item_Err(text)

        match alignment_hor:
            case 'center':
                alignment_hor = QtCore.Qt.AlignHCenter
            case 'left':
                alignment_hor = QtCore.Qt.AlignLeft
            case 'right':
                alignment_hor = QtCore.Qt.AlignRight

        item_font = QtGui.QFont(self.main.settings_custom['general']['ui_settings']['font_family'])
        item_font.setItalic(True)

        item.setFont(item_font)
        item.setTextAlignment(alignment_hor | QtCore.Qt.AlignVCenter)

        self.model().setItem(row, col, item)

    def update_ranks(self):
        data_prev = ''
        rank_prev = 1
        rank_next = 1

        sort_section = self.horizontalHeader().sortIndicatorSection()
        sort_order = self.horizontalHeader().sortIndicatorOrder()

        col_rank = self.find_header_hor(_tr('wl_tables', 'Rank'))

        self.sortByColumn(sort_section, sort_order)

        if sort_section != col_rank:
            self.disable_updates()

            for row in range(self.model().rowCount()):
                if not self.isRowHidden(row):
                    data_cur = self.model().item(row, sort_section).read_data()

                    if self.main.settings_custom['tables']['rank_settings']['continue_numbering_after_ties']:
                        if data_cur == data_prev:
                            self.model().item(row, col_rank).val = rank_prev
                            self.model().item(row, col_rank).setText(str(rank_prev))
                        else:
                            self.model().item(row, col_rank).val = rank_next
                            self.model().item(row, col_rank).setText(str(rank_next))

                            rank_prev = rank_next
                            rank_next += 1

                        data_prev = data_cur
                    else:
                        if data_cur == data_prev:
                            self.model().item(row, col_rank).val = rank_prev
                            self.model().item(row, col_rank).setText(str(rank_prev))
                        else:
                            self.model().item(row, col_rank).val = rank_next
                            self.model().item(row, col_rank).setText(str(rank_next))

                            rank_prev = rank_next

                        rank_next += 1
                        data_prev = data_cur

            self.enable_updates()

    def toggle_pct_data(self):
        self.disable_updates()

        match self.header_orientation:
            case 'hor':
                if self.table_settings['show_pct_data']:
                    for col in self.headers_pct:
                        if (
                            col not in self.cols_breakdown_file
                            or self.table_settings['show_breakdown_file']
                        ):
                            self.showColumn(col)
                else:
                    for col in self.headers_pct:
                        self.hideColumn(col)
            case 'vert':
                if self.table_settings['show_pct_data']:
                    for row in self.headers_pct:
                        self.showRow(row)
                else:
                    for row in self.headers_pct:
                        self.hideRow(row)

        self.enable_updates()

    def toggle_pct_data_span_position(self):
        self.disable_updates()

        match self.header_orientation:
            case 'hor':
                if self.table_settings['show_pct_data']:
                    for col in self.headers_pct:
                        if (
                            (
                                col not in self.cols_breakdown_file
                                or self.table_settings['show_breakdown_file']
                            ) and (
                                col not in self.cols_breakdown_span_position
                                or self.table_settings['show_breakdown_span_position']
                            )
                        ):
                            self.showColumn(col)
                else:
                    for col in self.headers_pct:
                        self.hideColumn(col)
            case 'vert':
                if self.table_settings['show_pct_data']:
                    for row in self.headers_pct:
                        self.showRow(row)
                else:
                    for row in self.headers_pct:
                        self.hideRow(row)

        self.enable_updates()

    def toggle_cum_data(self):
        precision_decimals = self.main.settings_custom['tables']['precision_settings']['precision_decimals']
        precision_pcts = self.main.settings_custom['tables']['precision_settings']['precision_pcts']

        # Boost performance
        if self.enable_sorting:
            self.sortByColumn(self.horizontalHeader().sortIndicatorSection(), self.horizontalHeader().sortIndicatorOrder())

        self.disable_updates()
        self.setSortingEnabled(False)

        match self.header_orientation:
            case 'hor':
                if self.table_settings['show_cum_data']:
                    for col in self.headers_cum:
                        val_cum = 0

                        # Integers
                        if col in self.headers_int:
                            for row in range(self.model().rowCount()):
                                if not self.isRowHidden(row):
                                    item = self.model().item(row, col)

                                    val_cum += item.val
                                    item.setText(str(val_cum))
                        # Floats
                        elif col in self.headers_float:
                            for row in range(self.model().rowCount()):
                                if not self.isRowHidden(row):
                                    item = self.model().item(row, col)

                                    val_cum += item.val
                                    item.setText(f'{val_cum:.{precision_decimals}}')
                        # Percentages
                        elif col in self.headers_pct:
                            for row in range(self.model().rowCount()):
                                if not self.isRowHidden(row):
                                    item = self.model().item(row, col)

                                    val_cum += item.val
                                    item.setText(f'{val_cum:.{precision_pcts}%}')
                else:
                    for col in self.headers_cum:
                        # Integers
                        if col in self.headers_int:
                            for row in range(self.model().rowCount()):
                                if not self.isRowHidden(row):
                                    item = self.model().item(row, col)

                                    item.setText(str(item.val))
                        # Floats
                        elif col in self.headers_float:
                            for row in range(self.model().rowCount()):
                                if not self.isRowHidden(row):
                                    item = self.model().item(row, col)

                                    item.setText(f'{item.val:.{precision_decimals}}')
                        # Percentages
                        elif col in self.headers_pct:
                            for row in range(self.model().rowCount()):
                                if not self.isRowHidden(row):
                                    item = self.model().item(row, col)

                                    item.setText(f'{item.val:.{precision_pcts}%}')
            case 'vert':
                if self.table_settings['show_cum_data']:
                    for row in self.headers_cum:
                        val_cum = 0

                        for col in range(self.model().columnCount() - 1):
                            item = self.model().item(row, col)

                            if not self.isColumnHidden(col) and not isinstance(item, Wl_Table_Item_Err):
                                val_cum += item.val

                                # Integers
                                if row in self.headers_int:
                                    item.setText(str(val_cum))
                                # Floats
                                elif row in self.headers_float:
                                    item.setText(f'{val_cum:.{precision_decimals}}')
                                # Percentages
                                elif row in self.headers_pct:
                                    item.setText(f'{val_cum:.{precision_pcts}%}')
                else:
                    for row in self.headers_cum:
                        for col in range(self.model().columnCount() - 1):
                            item = self.model().item(row, col)

                            if not self.isColumnHidden(col) and not isinstance(item, Wl_Table_Item_Err):
                                # Integers
                                if row in self.headers_int:
                                    item.setText(str(item.val))
                                # Floats
                                elif row in self.headers_float:
                                    item.setText(f'{item.val:.{precision_decimals}}')
                                # Percentages
                                elif row in self.headers_pct:
                                    item.setText(f'{item.val:.{precision_pcts}%}')

        self.enable_updates()

        if self.enable_sorting:
            self.setSortingEnabled(True)

    def toggle_breakdown_file(self):
        self.disable_updates()

        if self.table_settings['show_breakdown_file']:
            for col in self.cols_breakdown_file:
                if (
                    self.header_orientation == 'vert'
                    or col not in self.headers_pct
                    or self.table_settings['show_pct_data']
                ):
                    self.showColumn(col)
        else:
            for col in self.cols_breakdown_file:
                self.hideColumn(col)

        self.enable_updates()

    def toggle_breakdown_file_span_position(self):
        self.disable_updates()

        if self.table_settings['show_breakdown_file']:
            for col in self.cols_breakdown_file:
                if (
                    (
                        self.header_orientation == 'vert'
                        or col not in self.headers_pct
                        or self.table_settings['show_pct_data']
                    ) and (
                        col not in self.cols_breakdown_span_position
                        or self.table_settings['show_breakdown_span_position']
                    )
                ):
                    self.showColumn(col)
        else:
            for col in self.cols_breakdown_file:
                self.hideColumn(col)

        self.enable_updates()

    def toggle_breakdown_span_position(self):
        self.disable_updates()

        if self.table_settings['show_breakdown_span_position']:
            for col in self.cols_breakdown_span_position:
                if (
                    (
                        self.header_orientation == 'vert'
                        or col not in self.headers_pct
                        or self.table_settings['show_pct_data']
                    ) and (
                        col not in self.cols_breakdown_file
                        or self.table_settings['show_breakdown_file']
                    )
                ):
                    self.showColumn(col)
        else:
            for col in self.cols_breakdown_span_position:
                self.hideColumn(col)

        self.enable_updates()

    def filter_table(self):
        self.disable_updates()

        for i, row_filter in enumerate(self.row_filters):
            if row_filter:
                self.showRow(i)
            else:
                self.hideRow(i)

        self.enable_updates()

        if _tr('wl_tables', 'Rank') in self.get_header_labels_hor():
            self.update_ranks()

        if self.table_settings['show_cum_data']:
            self.toggle_cum_data()

    def generate_table(self):
        pass

    def generate_fig(self):
        pass

    def clr_table(self, num_headers = 1, confirm = False):
        confirmed = True

        # Ask for confirmation if results have not been exported
        if confirm:
            if not self.is_empty() and not self.results_saved:
                confirmed = wl_dialogs.Wl_Dialog_Question(
                    self.main,
                    title = _tr('wl_tables', 'Clear Table'),
                    text = _tr('wl_tables', '''
                        <div>The results in the table have yet been exported.</div>
                        <br>
                        <div>Do you want to clear all results in the table?</div>
                    ''')
                ).exec()

        if confirmed:
            self.model().clear()

            match self.header_orientation:
                case 'hor':
                    self.horizontalHeader().blockSignals(True)

                    self.model().setColumnCount(len(self.headers))
                    self.model().setRowCount(num_headers)

                    self.model().setHorizontalHeaderLabels(self.headers)

                    self.horizontalHeader().blockSignals(False)

                    self.horizontalHeader().sectionCountChanged.emit(0, num_headers)
                case 'vert':
                    self.verticalHeader().blockSignals(True)

                    self.model().setRowCount(len(self.headers))
                    self.model().setColumnCount(num_headers)

                    self.model().setVerticalHeaderLabels(self.headers)

                    self.verticalHeader().blockSignals(False)

                    self.verticalHeader().sectionCountChanged.emit(0, num_headers)

            for i in range(self.model().rowCount()):
                self.showRow(i)

            for i in range(self.model().columnCount()):
                self.showColumn(i)

            self.headers_int = {self.find_header(header) for header in self.headers_int_old}
            self.headers_float = {self.find_header(header) for header in self.headers_float_old}
            self.headers_pct = {self.find_header(header) for header in self.headers_pct_old}
            self.headers_cum = {self.find_header(header) for header in self.headers_cum_old}

            self.cols_breakdown_file = {
                self.find_header_hor(col)
                for col in self.cols_breakdown_file_old
            }
            self.cols_breakdown_span_position = {
                self.find_header_hor(col)
                for col in self.cols_breakdown_span_position_old
            }

            self.results_saved = False

            self.model().itemChanged.emit(QtGui.QStandardItem())

        return confirmed

# Avoid circular imports
from wordless.wl_results import wl_results_filter, wl_results_search, wl_results_sort # pylint: disable=wrong-import-position

class Wl_Table_Data_Search(Wl_Table_Data):
    def __init__(
        self, main, tab,
        headers, header_orientation = 'hor',
        headers_int = None, headers_float = None,
        headers_pct = None, headers_cum = None,
        cols_breakdown_file = None, cols_breakdown_span_position = None,
        enable_sorting = False, generate_fig = True
    ):
        super().__init__(
            main, tab,
            headers, header_orientation,
            headers_int, headers_float,
            headers_pct, headers_cum,
            cols_breakdown_file, cols_breakdown_span_position,
            enable_sorting, generate_fig
        )

        self.model().itemChanged.connect(self.results_changed)

        self.label_num_results = QtWidgets.QLabel('', self)
        self.button_results_search = wl_buttons.Wl_Button(_tr('wl_tables', 'Search in results'), self)
        self.dialog_results_search = wl_results_search.Wl_Dialog_Results_Search(self.main, table = self)

        self.button_results_search.setMinimumWidth(140)

        self.button_generate_table.clicked.connect(self.dialog_results_search.clr_history)
        self.button_results_search.clicked.connect(self.dialog_results_search.show)

        self.results_changed()

    def results_changed(self):
        rows_visible = len([i for i in range(self.model().rowCount()) if not self.isRowHidden(i)])

        if not self.is_empty() and rows_visible:
            self.label_num_results.setText(_tr('wl_tables', 'Number of results: ') + str(rows_visible))

            self.button_results_search.setEnabled(True)
        else:
            self.label_num_results.setText(_tr('wl_tables', 'Number of results: 0'))

            self.button_results_search.setEnabled(False)

        self.results_changed_menu_edit()

    def results_changed_menu_edit(self):
        if self.button_results_search.isEnabled():
            self.main.action_edit_results_search.setEnabled(True)
        else:
            self.main.action_edit_results_search.setEnabled(False)

        self.main.action_edit_results_filter.setEnabled(False)
        self.main.action_edit_results_sort.setEnabled(False)

class Wl_Table_Data_Sort_Search(Wl_Table_Data):
    def __init__(
        self, main, tab,
        headers, header_orientation = 'hor',
        headers_int = None, headers_float = None,
        headers_pct = None, headers_cum = None,
        cols_breakdown_file = None, cols_breakdown_span_position = None,
        enable_sorting = False, generate_fig = True
    ):
        super().__init__(
            main, tab,
            headers, header_orientation,
            headers_int, headers_float,
            headers_pct, headers_cum,
            cols_breakdown_file, cols_breakdown_span_position,
            enable_sorting, generate_fig
        )

        self.model().itemChanged.connect(self.results_changed)

        self.label_num_results = QtWidgets.QLabel('', self)
        self.button_results_sort = wl_buttons.Wl_Button(_tr('wl_tables', 'Sort results'), self)
        self.button_results_search = wl_buttons.Wl_Button(_tr('wl_tables', 'Search in results'), self)

        self.dialog_results_sort = wl_results_sort.Wl_Dialog_Results_Sort_Concordancer(self.main, table = self)
        self.dialog_results_search = wl_results_search.Wl_Dialog_Results_Search(self.main, table = self)

        self.button_results_sort.setMinimumWidth(140)
        self.button_results_search.setMinimumWidth(140)

        self.button_generate_table.clicked.connect(self.dialog_results_search.clr_history)
        self.button_results_sort.clicked.connect(self.dialog_results_sort.show)
        self.button_results_search.clicked.connect(self.dialog_results_search.show)

        self.results_changed()

    def results_changed(self):
        rows_visible = len([i for i in range(self.model().rowCount()) if not self.isRowHidden(i)])

        if not self.is_empty() and rows_visible:
            self.label_num_results.setText(_tr('wl_tables', 'Number of results: ') + str(rows_visible))

            self.button_results_sort.setEnabled(True)
            self.button_results_search.setEnabled(True)
        else:
            self.label_num_results.setText(_tr('wl_tables', 'Number of results: 0'))

            self.button_results_sort.setEnabled(False)
            self.button_results_search.setEnabled(False)

        self.results_changed_menu_edit()

    def results_changed_menu_edit(self):
        if self.button_results_search.isEnabled():
            self.main.action_edit_results_search.setEnabled(True)
        else:
            self.main.action_edit_results_search.setEnabled(False)

        if self.button_results_sort.isEnabled():
            self.main.action_edit_results_sort.setEnabled(True)
        else:
            self.main.action_edit_results_sort.setEnabled(False)

        self.main.action_edit_results_filter.setEnabled(False)

class Wl_Table_Data_Filter_Search(Wl_Table_Data):
    def __init__(
        self, main, tab,
        headers, header_orientation = 'hor',
        headers_int = None, headers_float = None,
        headers_pct = None, headers_cum = None,
        cols_breakdown_file = None, cols_breakdown_span_position = None,
        enable_sorting = False, generate_fig = True
    ):
        super().__init__(
            main, tab,
            headers, header_orientation,
            headers_int, headers_float,
            headers_pct, headers_cum,
            cols_breakdown_file, cols_breakdown_span_position,
            enable_sorting, generate_fig
        )

        self.model().itemChanged.connect(self.results_changed)

        self.label_num_results = QtWidgets.QLabel('', self)
        self.button_results_filter = wl_buttons.Wl_Button(_tr('wl_tables', 'Filter results'), self)
        self.button_results_search = wl_buttons.Wl_Button(_tr('wl_tables', 'Search in results'), self)

        self.dialog_results_search = wl_results_search.Wl_Dialog_Results_Search(self.main, table = self)

        self.button_results_filter.setMinimumWidth(140)
        self.button_results_search.setMinimumWidth(140)

        self.button_generate_table.clicked.connect(self.dialog_results_search.clr_history)
        self.button_results_filter.clicked.connect(self.results_filter_clicked)
        self.button_results_search.clicked.connect(self.dialog_results_search.show)

        self.results_changed()

    def results_changed(self):
        rows_visible = len([i for i in range(self.model().rowCount()) if not self.isRowHidden(i)])

        if not self.is_empty():
            self.label_num_results.setText(_tr('wl_tables', 'Number of results: ') + str(rows_visible))

            self.button_results_filter.setEnabled(True)
        else:
            self.label_num_results.setText(_tr('wl_tables', 'Number of results: 0'))

            self.button_results_filter.setEnabled(False)

        if not self.is_empty() and rows_visible:
            self.button_results_search.setEnabled(True)
        else:
            self.button_results_search.setEnabled(False)

        self.results_changed_menu_edit()

    def results_changed_menu_edit(self):
        if self.button_results_search.isEnabled():
            self.main.action_edit_results_search.setEnabled(True)
        else:
            self.main.action_edit_results_search.setEnabled(False)

        if self.button_results_filter.isEnabled():
            self.main.action_edit_results_filter.setEnabled(True)
        else:
            self.main.action_edit_results_filter.setEnabled(False)

        self.main.action_edit_results_sort.setEnabled(False)

    def results_filter_clicked(self):
        match self.tab:
            case 'dependency_parser':
                wl_dialog_results_filter = wl_results_filter.Wl_Dialog_Results_Filter_Dependency_Parser(
                    self.main,
                    table = self
                )
            case 'wordlist_generator' | 'ngram_generator':
                wl_dialog_results_filter = wl_results_filter.Wl_Dialog_Results_Filter_Wordlist_Generator(
                    self.main,
                    table = self
                )
            case 'collocation_extractor' | 'colligation_extractor' | 'keyword_extractor':
                wl_dialog_results_filter = wl_results_filter.Wl_Dialog_Results_Filter_Collocation_Extractor(
                    self.main,
                    table = self
                )

        wl_dialog_results_filter.show()
